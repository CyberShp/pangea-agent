from __future__ import annotations

import atexit
import json
import logging
import shutil
import subprocess
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from types import SimpleNamespace
from typing import Callable

from docx import Document
from PIL import Image
from openpyxl import Workbook

from pangea_agent.agent_io import agent_path, read_json, write_json
from pangea_agent.cli.run_module_analysis import (
    _reserve_run_id,
    resume_module_analysis,
    run_module_analysis,
    start_module_analysis,
)
from pangea_agent.cli.adapter_api import (
    bind_action,
    settle_action,
    validate_action,
)
from pangea_agent.documents.coverage import match_coverage_records, parse_coverage_xlsx
from pangea_agent.graph.nodes.load_contract import load_contract
from pangea_agent.graph.nodes.resolve_repositories import resolve_repositories
from pangea_agent.graph.nodes.locate_module import locate_module
from pangea_agent.graph.nodes.index_materials import index_materials
from pangea_agent.graph.nodes.make_analysis_units import _cluster_groups
from pangea_agent.graph.nodes.prepare_worker_tasks import (
    _coverage_context,
    _related_state_context,
    _source_inventory,
)
from pangea_agent.graph.semantic_checks import build_runtime_semantic_checks
from pangea_agent.graph.run_store import load_worker_task, review_result_skeleton
from pangea_agent.graph.validation import (
    ArtifactRejected,
    _finding_excludes_linked_leak,
    _known_c_container_semantic_artifact_ids,
    _known_local_allocation_concurrency_artifact_ids,
    _known_c_precheck_order_artifact_ids,
    _known_c_review_misread_finding_ids,
    _known_c_unmap_finding_ids,
    _known_retained_realloc_artifact_ids,
    _validate_known_c_required_failure_paths,
    _non_actionable_review_issue_ids,
    _retained_realloc_source_paths,
    _reviewer_owned_field_issue_ids,
    _reviewer_self_correction_issue_ids,
    _stale_artifact_restoration_issue_ids,
)
from pangea_agent.models.worker import (
    AnalysisUnit,
    IndependentReviewResult,
    ReviewResult,
    ReviewTask,
    WorkerResult,
    WorkerTask,
)
from pangea_agent.models.contract import TaskContract
from pangea_agent.inventory.source_scanner import _known_macro_parse_artifact
from pangea_agent.inventory.lua_symbols import parse_lua_file
from pangea_agent.inventory.scope_expander import MAX_DIRECT_CALLERS_PER_GROUP
from pangea_agent.inventory.source_languages import checkpoint_rubrics
from pangea_agent.report.html import render_html_report
from pangea_agent.report.markdown import render_report


Scenario = Callable[[], None]
_TEMP_ROOTS: list[Path] = []
_ANALYSIS_SESSION_ID = "00000000-0000-4000-8000-000000000001"
_REPLACEMENT_SESSION_ID = "00000000-0000-4000-8000-000000000002"
_SECOND_REPLACEMENT_SESSION_ID = "00000000-0000-4000-8000-000000000003"
_REVIEW_SESSION_ID = "00000000-0000-4000-8000-000000000011"


@atexit.register
def _cleanup_temp_roots() -> None:
    for root in _TEMP_ROOTS:
        shutil.rmtree(root, ignore_errors=True)


def _workspace(*, run_id: str = "smoke-01", repositories: tuple[str, ...] = ("demo",)) -> tuple[Path, Path, Path]:
    root = Path(tempfile.mkdtemp(prefix="pangea-v1-smoke-"))
    _TEMP_ROOTS.append(root)
    data_root = root / "data"
    for repo_id in repositories:
        repo = data_root / "repositories" / repo_id
        module = repo / "module"
        module.mkdir(parents=True)
        (module / "entry.c").write_text(f"int {repo_id}_start(void) {{ return 0; }}\n", encoding="utf-8")
        subprocess.run(["git", "init", "-q", str(repo)], check=True)
        subprocess.run(["git", "-C", str(repo), "add", "module/entry.c"], check=True)
        subprocess.run(
            [
                "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
                "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm", "fixture",
            ],
            check=True,
        )
    contract = {
        "run_id": run_id,
        "data_root": str(data_root),
        "mode": "module_analysis",
        "repositories": list(repositories),
        "target": "CHAP",
        "source_scope": ["module"],
    }
    contract_path = root / "contract.json"
    write_json(contract_path, contract)
    return root, data_root, contract_path


def _lua_workspace() -> tuple[Path, Path, Path]:
    root, data_root, contract_path = _workspace(repositories=("lua-demo",))
    repo = data_root / "repositories" / "lua-demo"
    module = repo / "module"
    (module / "entry.c").unlink()
    (module / "helper.lua").write_text(
        "local M = {}\nfunction M.value() return 1 end\nreturn M\n",
        encoding="utf-8",
    )
    (module / "component.lua").write_text(
        """local mc = require("mc")
local helper = require("module.helper")
local Component = mc.class("Component")
Component.changed = mc.signal()

function Component:ctor()
    self.ready = false
end

function Component:pre_init()
    self.value = helper.value()
end

function Component:init()
    self.changed:connect(function() self.ready = true end)
end

function Component:run()
    if not self.ready then error("not ready") end
    return pcall(function() self.changed:emit() end)
end

return Component
""",
        encoding="utf-8",
    )
    subprocess.run(["git", "-C", str(repo), "add", "-A"], check=True)
    subprocess.run(
        [
            "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
            "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm", "lua fixture",
        ],
        check=True,
    )
    contract = read_json(contract_path)
    contract["source_scope"] = ["module/component.lua"]
    write_json(contract_path, contract)
    return root, data_root, contract_path


def _task_result(task: dict, *, finish_reason: str = "stop", fake_location: bool = False) -> dict:
    repo_id = task["unit"]["repo_id"]
    source_path = task["unit"]["source_scope"][0]
    chunk_id = f"{repo_id}:{source_path}:1-1"
    location = "fake:missing.c:1-1" if fake_location else chunk_id
    evidence = {"chunk_id": chunk_id, "location": location, "observation": "模块入口实现"}
    stage = task["stage"]
    result = {
        "schema_version": "1.0",
        "run_id": task["run_id"],
        "unit_id": task["unit"]["unit_id"],
        "worker_id": f"worker-{task['unit']['unit_id']}",
        "attempt": task["attempt"],
        "completed_stage": stage,
        "finish_reason": finish_reason,
        "summary": "完成当前单元分析",
        "analyzed_scope": task["unit"]["source_scope"],
        "analyzed_context_scope": task["unit"].get("context_scope", []),
        "evidence": [evidence],
        "business_flows": [{
            "title": "模块启动",
            "description": "测试人员启动目标模块。",
            "steps": ["启动目标模块"],
            "mermaid": None,
            "evidence": [evidence],
        }],
        "visual_findings": [],
        "risks": [],
        "test_cases": [],
        "addressed_review_issue_ids": [item["issue_id"] for item in task["review_issues"]],
        "errors": [] if finish_reason == "stop" else [finish_reason],
        "analysis_checkpoint": {
            "source_paths_reviewed": task["unit"]["source_scope"],
            "lifecycle_stages_checked": ["初始化", "运行", "停止", "恢复"],
            "failure_paths": [{
                "path_id": item["check_id"],
                "linked_risk_ids": [],
                "trigger": "入口调用",
                "side_effects": "进入模块逻辑",
                "failure": "无已确认故障",
                "caller_handling": "调用方读取返回值",
                "final_states": "模块保持可用",
                "disposition": "excluded",
            } for item in task.get("semantic_check_items", [])] or [{
                "path_id": "F-001",
                "linked_risk_ids": [],
                "trigger": "入口调用",
                "side_effects": "进入模块逻辑",
                "failure": "无已确认故障",
                "caller_handling": "调用方读取返回值",
                "final_states": "模块保持可用",
                "disposition": "excluded",
            }],
            "material_decisions": [],
            "coverage_priorities": [],
            "coverage_decisions": [],
            "risk_set_frozen": True,
            "counterexamples_checked": ["异常返回不会被误写为成功"],
        },
    }
    if stage == "source_checkpoint":
        result["evidence"] = []
        result["business_flows"] = []
        result["analysis_checkpoint"]["risk_set_frozen"] = False
        result["analysis_checkpoint"]["counterexamples_checked"] = []
    elif stage == "risk_analysis":
        result["test_cases"] = []
        result["analysis_checkpoint"]["counterexamples_checked"] = []
    for failure_path in result["analysis_checkpoint"]["failure_paths"]:
        if failure_path["path_id"].endswith(":retry"):
            failure_path["trigger"] = "首次失败后修正条件并重试成功，再 emit/update 一次 normal 事件"
            failure_path["side_effects"] = "按实际注册顺序执行本次事件"
            failure_path["final_states"] = "单次 emit/update 后状态已记录"
        elif failure_path["path_id"].endswith(":multi-instance"):
            failure_path["trigger"] = "同一运行时先由 A emit normal，再由 B emit normal"
            failure_path["side_effects"] = "两次 normal 均按注册顺序执行"
            failure_path["final_states"] = "A normal 后 B normal 的最终状态已记录"
    if stage != "source_checkpoint":
        manifest = read_json(Path(task["source_manifest_path"]))
        result["analysis_checkpoint"]["material_decisions"] = [
            {
                "path": item["path"],
                "decision": "context",
                "reason": "smoke 默认将已解析资料保留为上下文；资料专项场景会覆盖该决定。",
            }
            for item in manifest.get("material_catalog", [])
            if item.get("type") == "material"
            and str(item.get("parse_status", "")).startswith("parsed")
        ]
    return result


def _validate_worker_task(task_path: Path) -> None:
    subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result", "--task", str(task_path)],
        check=True,
        capture_output=True,
    )


def _reject_worker_task(task_path: Path) -> str:
    rejected = subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result", "--task", str(task_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    return rejected.stderr


def _validate_review_task(task_path: Path) -> None:
    subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "check-review-artifact", "--task", str(task_path)],
        check=True,
        capture_output=True,
    )


def _reject_review_task(task_path: Path) -> str:
    rejected = subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "check-review-artifact", "--task", str(task_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    return rejected.stderr


def _bind_analysis_actions(state: dict) -> None:
    for index, action in enumerate(state.get("agent_actions", [])):
        if action.get("role") != "analysis" or action.get("task_id") is not None:
            continue
        task_path = Path(action["task_path"])
        run_dir = task_path.parents[2]
        session_id = f"10000000-0000-4000-8000-{index + 1:012d}"
        subprocess.run(
            [
                sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
                "--run-id", run_dir.name, "--data-root", str(run_dir.parent.parent),
                "--role", "analysis", "--unit-id", action["unit_id"], "--task-id", session_id,
            ],
            check=True,
            capture_output=True,
        )


def _advance_to_test_generation(state: dict) -> dict:
    while state["phase"] in {"WAITING_SOURCE_CHECKPOINT", "WAITING_RISK_ANALYSIS"}:
        _bind_analysis_actions(state)
        task_paths = [Path(path) for path in state["agent_task_paths"]]
        for task_path in task_paths:
            task = read_json(task_path)
            result = _task_result(task)
            write_json(Path(task["result_path"]), result)
            _validate_worker_task(task_path)
        run_dir = task_paths[0].parents[2]
        state = run_module_analysis(str(run_dir / "inputs" / "task-contract.json"))
    assert state["phase"] == "WAITING_TEST_GENERATION"
    return state


def _test_case(
    case_id: str,
    *,
    risk_ids: list[str] | None = None,
    requirement_ids: list[str] | None = None,
    material_ids: list[str] | None = None,
    coverage_ids: list[str] | None = None,
    title: str = "业务行为验证",
) -> dict:
    step = {
        "action": "执行目标业务操作",
        "expected_result": "系统表现符合测试依据",
    }
    if risk_ids:
        step["failure_observation"] = "系统出现关联风险描述的错误业务表现"
    return {
        "test_case_id": case_id,
        "title": title,
        "case_type": "功能",
        "linked_risk_ids": risk_ids or [],
        "linked_requirement_ids": requirement_ids or [],
        "linked_material_ids": material_ids or [],
        "linked_coverage_ids": coverage_ids or [],
        "preconditions": ["环境就绪"],
        "steps": [step],
        "observability": ["业务结果"],
        "cleanup": ["恢复环境"],
    }


def _unpaired_test_step_rejected() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["risks"] = [{
        "risk_id": "U00-R1",
        "title": "可执行风险",
        "affected_paths": task["unit"]["source_scope"],
        "dfx": ["功能与状态"],
        "severity": "Medium",
        "confidence": "high",
        "trigger": "业务触发",
        "system_result": "系统异常",
        "external_observation": "外部可观测",
        "exclusion_condition": "排除条件",
        "upstream_semantics": {
            "reachability": "业务入口可达",
            "caller_constraints": "调用方未消除",
            "documented_behavior": "资料未定义为预期",
            "existing_tests": "已有测试未覆盖",
            "conclusion": "risk_remains",
        },
        "translation_status": "Blackbox-ready",
        "status": "pending",
        "evidence": result["evidence"],
    }]
    result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R1"],
        "failure": "业务操作失败",
        "final_states": "系统状态异常",
        "disposition": "risk",
    })
    result["test_cases"] = [{
        "test_case_id": "U00-TC1",
        "title": "步骤与预期错位",
        "case_type": "功能",
        "linked_risk_ids": ["U00-R1"],
        "linked_requirement_ids": [],
        "linked_material_ids": [],
        "linked_coverage_ids": [],
        "preconditions": ["环境就绪"],
        "steps": [{"action": "准备环境"}],
        "observability": ["外部日志"],
        "cleanup": ["恢复环境"],
    }]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "expected_result" in rejection
    result["test_cases"][0]["steps"] = [{
        "action": "准备环境",
        "expected_result": "环境准备成功",
    }]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "failure_observation" in rejection
    result["test_cases"][0]["steps"] = [{
        "action": "检查 B.audit_count",
        "expected_result": "B.audit_count=0（通过标准）",
        "failure_observation": "B.audit_count=0（命中即 FAIL）",
    }]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "PASS 与 FAIL 状态断言相同" in rejection


def _paired_test_steps_render_together() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    state["test_cases"] = [_test_case("TC-PAIR", risk_ids=["R-PAIR-01"])]
    markdown = render_report(state)
    html = render_html_report(state)
    assert "**操作目标**：执行目标业务操作" in markdown
    assert "**通过标准**：系统表现符合测试依据" in markdown
    assert "**已知失败观测（命中即 FAIL）**：系统出现关联风险描述的错误业务表现" in markdown
    assert "<td>执行目标业务操作</td><td>系统表现符合测试依据</td><td>系统出现关联风险描述的错误业务表现</td>" in html


def _expected_result_cannot_include_current_behavior() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    case = _test_case("TC-ORACLE", requirement_ids=["REQ-ORACLE-01"])
    case["steps"][0]["expected_result"] = (
        "本次 callback_count 增量为 1（当前实现错误地增加 2）"
    )
    second = _test_case("TC-ORACLE-SECOND", requirement_ids=["REQ-ORACLE-02"])
    second["steps"][0]["expected_result"] = "源码当前会返回失败，实测即复现风险"
    third = _test_case("TC-ORACLE-CONTRAST", requirement_ids=["REQ-ORACLE-03"])
    third["steps"][0]["expected_result"] = (
        "B.callback_count 也变为 1（正确预期是 B.callback_count 保持为 0）"
    )
    result["test_cases"] = [case, second, third]
    write_json(Path(task["result_path"]), result)
    reason = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "expected_result 混入当前实现行为" in reason
    assert "TC-ORACLE 第 1 步" in reason
    assert "TC-ORACLE-SECOND 第 1 步" in reason
    assert "TC-ORACLE-CONTRAST 第 1 步" in reason


def _requirement_only_test_case_is_accepted() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["test_cases"] = [_test_case(
        "U00-REQ-TC1",
        requirement_ids=["REQ-DEMO-01"],
        title="需求行为补测",
    )]
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"


def _semantic_check_risk_scope_is_enforced() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    task["semantic_check_items"] = [{
        "check_id": "SC-DEMO-PAIR-01",
        "kind": "paired_operation",
        "subject_path": "module/entry.c",
        "instruction": "只检查当前实现",
        "context_paths": ["module/entry.c"],
    }]
    write_json(task_path, task)
    result = _task_result(task)
    result["risks"] = [{
        "risk_id": "U00-R-SCOPE",
        "title": "错误实现范围",
        "affected_paths": ["module/other.c"],
        "dfx": ["功能与状态"],
        "severity": "Medium",
        "confidence": "high",
        "trigger": "业务触发",
        "system_result": "系统异常",
        "external_observation": "外部可观测",
        "exclusion_condition": "排除条件",
        "upstream_semantics": {
            "reachability": "业务入口可达",
            "caller_constraints": "调用方未消除",
            "documented_behavior": "资料未定义为预期",
            "existing_tests": "已有测试未覆盖",
            "conclusion": "risk_remains",
        },
        "translation_status": "Developer-confirm",
        "status": "pending",
        "evidence": result["evidence"],
    }]
    result["analysis_checkpoint"]["failure_paths"] = [{
        "path_id": "SC-DEMO-PAIR-01",
        "linked_risk_ids": ["U00-R-SCOPE"],
        "trigger": "入口调用",
        "side_effects": "进入模块逻辑",
        "failure": "配对失败",
        "caller_handling": "调用方继续",
        "final_states": "系统异常",
        "disposition": "risk",
    }]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(task_path)
    assert "未落在当前分析单元 source_scope" in rejection

    task["semantic_check_items"][0]["subject_path"] = "module/different.c"
    write_json(task_path, task)
    result["risks"][0]["affected_paths"] = ["demo:module/entry.c:1-8"]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(task_path)
    assert "未声明 semantic check" in rejection

    task["semantic_check_items"][0]["subject_path"] = "module/entry.c"
    write_json(task_path, task)
    result["risks"][0]["affected_paths"] = ["demo:module/entry.c:1-8 (paired operation)"]
    result["analysis_checkpoint"]["failure_paths"][0]["linked_risk_ids"] = []
    result["analysis_checkpoint"]["failure_paths"][0]["disposition"] = "excluded"
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(task_path)
    assert "没有被当前 risk/unresolved failure path 关联" in rejection

    result["analysis_checkpoint"]["failure_paths"][0]["linked_risk_ids"] = [
        "U00-R-SCOPE"
    ]
    result["analysis_checkpoint"]["failure_paths"][0]["disposition"] = "risk"
    result["analysis_checkpoint"]["failure_paths"].append({
        "path_id": "FP-UNLINKED-RISK",
        "linked_risk_ids": [],
        "trigger": "另一业务入口调用",
        "side_effects": "已发生状态变化",
        "failure": "已确认的故障",
        "caller_handling": "调用方未清理",
        "final_states": "模块状态异常",
        "disposition": "risk",
    })
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(task_path)
    assert "risk/unresolved failure path 尚未转换并关联 RiskCard" in rejection

    result["analysis_checkpoint"]["failure_paths"].pop()
    original_system_result = result["risks"][0]["system_result"]
    result["risks"][0]["system_result"] = (
        "TAILQ_REMOVE silently succeeds as a no-op for an element outside the queue"
    )
    assert _known_c_container_semantic_artifact_ids(
        WorkerTask.model_validate(task),
        WorkerResult.model_validate(result),
    ) == ["U00-R-SCOPE"]

    result["risks"][0]["system_result"] = original_system_result
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(task_path)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"

    frozen_source = Path(task["repositories"][0]["source_root"]) / "module/entry.c"
    frozen_source.write_text(
        """void *build(unsigned long size) {
    void *page;
    page = calloc(1, size);
    void *new_page = realloc(page, size + 1);
    if (new_page == NULL) {
        break;
    }
    page = new_page;
    page_count++;
    return page;
}
""",
        encoding="utf-8",
    )
    worker_task = WorkerTask.model_validate(task)
    assert _retained_realloc_source_paths(worker_task) == {"module/entry.c"}
    realloc_result = WorkerResult.model_validate(result)
    realloc_result.risks[0].system_result = (
        "realloc failure breaks without free and leaks the original memory"
    )
    assert _known_retained_realloc_artifact_ids(worker_task, realloc_result) == [
        "U00-R-SCOPE"
    ]
    realloc_result.risks[0].system_result = (
        "When realloc fails, page_count was already incremented and becomes inconsistent"
    )
    assert _known_retained_realloc_artifact_ids(worker_task, realloc_result) == [
        "U00-R-SCOPE"
    ]
    realloc_result.risks[0].system_result = (
        "Two concurrent calls share page; another call frees page and leaves a stale pointer"
    )
    assert _known_local_allocation_concurrency_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]

    frozen_source.write_text(
        """#define AE4DMA_PCIE_BAR 0
static inline void spdk_mmio_write_4(volatile uint32_t *dst, uint32_t val) {}
/** spdk_ae4dma_build_copy return 0 on success, negative errno on failure. */
int build_copy(void *src, void *dst) {
    psrc_addr = spdk_vtophys(src, &src_len);
    pdst_addr = spdk_vtophys(dst, &dst_len);
    if (psrc_addr == SPDK_VTOPHYS_ERROR || pdst_addr == SPDK_VTOPHYS_ERROR) {
        return -EFAULT;
    }
    seg_len = spdk_min(src_len, dst_len);
    if (seg_len == 0) {
        return -EINVAL;
    }
    if (cmd_q->ring_buff_count >= RING_LIMIT) {
        return 1;
    }
    return ae4dma_prep_copy(dst, src, seg_len);
}
int build_batch(void *src, void *dst) {
    if (spdk_vtophys(src, &src_len) == SPDK_VTOPHYS_ERROR) {
        return -EFAULT;
    }
    cb_desc = ae4dma_prep_copy(dst, src, seg_len);
    if (!cb_desc) {
        return -ENOMEM;
    }
    last_desc = cb_desc;
    if (last_desc) {
        cb_desc->callback_fn = cb_fn;
    }
    return 0;
}
int process_events(struct cmd_queue *cmd_q) {
    uint64_t sub_desc_cnt = cmd_q->ring_buff_count;
    while (sub_desc_cnt) {
    if (desc_status == AE4DMA_DMA_DESC_SUBMITTED) {
        break;
    }
    if (desc_status != AE4DMA_DMA_DESC_COMPLETED) {
        desc_err_code = hw_desc->err_code;
    }
    assert(cmd_q->ring_buff_count > 0);
    cmd_q->ring_buff_count--;
    if (ring[tail].callback_fn) {
        ring[tail].callback_fn(ring[tail].callback_arg, desc_err_code);
    }
    tail = (tail + 1);
    sub_desc_cnt--;
    }
}
/** Flush previously built descriptors and flush the descriptor to hardware for further processing. */
void spdk_ae4dma_flush(struct chan *ae4dma, int hwq_id) {}
static int ae4dma_map_pci_bar(struct chan *ae4dma) {
    return -1;
}
static int ae4dma_unmap_pci_bar(struct chan *ae4dma) {
    int rc = 0;
    rc = spdk_pci_device_unmap_bar(ae4dma->device, 0, addr);
    return rc;
}
static bool ae4dma_config_queues_per_device(uint8_t num_hw_queues) {
    if (num_hw_queues <= AE4DMA_MAX_HW_QUEUES) {
        return false;
    }
    return true;
}
static bool ae4dma_desc_cmdq_full(struct cmd_queue *cmd_q) {
    return cmd_q->count >= (AE4DMA_DESCRIPTORS_PER_CMDQ - 4);
}
static void ae4dma_prep_copy(struct cmd_queue *cmd_q) {
    if (cmd_q->ring_buff_count >= (AE4DMA_DESCRIPTORS_PER_CMDQ - 4)) {
        return;
    }
}
int ae4dma_channel_start(struct chan *ae4dma) {
    if (!ae4dma_config_queues_per_device(hw_queues)) {
        q_per_eng = hw_queues;
    } else {
        q_per_eng = AE4DMA_MAX_HW_QUEUES;
    }
    size = cmd_q->queue_size;
    cmd_q->qring_buffer_pa = spdk_vtophys(cmd_q->qbase_addr, &size);
    cmd_q->ring = calloc(RING_LIMIT, sizeof(*cmd_q->ring));
    if (!cmd_q->ring) return -ENOMEM;
}
void ae4dma_channel_destruct(struct chan *ae4dma) {
    spdk_pci_device_unmap_bar(ae4dma->device, 0, addr);
    ae4dma_unmap_pci_bar(ae4dma);
    spdk_free(ae4dma->cmd_q[0].qbase_addr);
    free(ae4dma->cmd_q[0].ring);
}
void spdk_ae4dma_detach(struct chan *ae4dma) {
    ae4dma_channel_destruct(ae4dma);
    free(ae4dma);
}
void *ae4dma_attach(struct chan *ae4dma) {
    if (ae4dma_channel_start(ae4dma) != 0) {
        ae4dma_channel_destruct(ae4dma);
        free(ae4dma);
        return NULL;
    }
}
""",
        encoding="utf-8",
    )
    realloc_result.risks[0].trigger = (
        "vtophys returns SPDK_VTOPHYS_ERROR; 当前描述符已入队写入 ring"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    required_result = WorkerResult.model_validate(result)
    try:
        _validate_known_c_required_failure_paths(worker_task, required_result)
    except ArtifactRejected as exc:
        assert "ring-full" in str(exc)
    else:
        raise AssertionError("ring-full positive return contract omission must be rejected")
    contract_path = required_result.analysis_checkpoint.failure_paths[0]
    contract_path.path_id = "prep_copy-ring-full"
    contract_path.disposition = "risk"
    contract_path.trigger = "capacity guard is reached"
    contract_path.failure = "returns positive 1 while public contract requires negative errno"
    contract_path.caller_handling = "caller observes the contract violation"
    try:
        _validate_known_c_required_failure_paths(worker_task, required_result)
    except ArtifactRejected as exc:
        assert "unmap" in str(exc)
    else:
        raise AssertionError("lost destructor unmap rc omission must be rejected")
    unmap_path = contract_path.model_copy(deep=True)
    unmap_path.path_id = "FP-UNMAP-RC"
    unmap_path.trigger = "detach calls destruct"
    unmap_path.failure = "destruct ignores the unmap rc return value"
    unmap_path.caller_handling = "detach drops the nonzero return value"
    required_result.analysis_checkpoint.failure_paths.append(unmap_path)
    _validate_known_c_required_failure_paths(worker_task, required_result)
    callback_path = contract_path.model_copy(deep=True)
    callback_path.path_id = "prep_copy-callback-wrong-desc"
    callback_path.trigger = "normal multi-segment completion"
    callback_path.failure = "functionally correct today but considered a coding-practice issue"
    callback_path.caller_handling = "verify the final callback placement"
    required_result.analysis_checkpoint.failure_paths.append(callback_path)
    assert "prep_copy-callback-wrong-desc" in _known_c_precheck_order_artifact_ids(
        worker_task,
        required_result,
    )
    required_result.analysis_checkpoint.failure_paths.pop()
    realloc_result.risks[0].trigger = (
        "vtophys returns SPDK_VTOPHYS_ERROR; 当前描述符未入队，"
        "此前成功段的描述符已入队且无法回滚"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == []
    realloc_result.risks[0].trigger = (
        "错误状态出现后 callback 未被调用，调用方一直等待"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "初始化返回 NULL，若调用方未检查就使用会触发空指针崩溃"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "vtophys 失败后 ring 中存在无效物理地址描述符，硬件将访问错误地址"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "seg_len=0 返回 -EINVAL，当前 segment 无描述符入队且 ring 无残留"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "ring full 返回 1；重新调用会覆盖未消费的旧描述符"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "返回值不含错误状态，但 callback 会收到 err_code"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "ring calloc 分配失败后 qbase_addr 未释放并发生资源泄漏"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "process_events 未检查 ring_buff_count 就执行 assert，计数下溢触发断言"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "spdk_vtophys 的 size 未初始化，长度参数错误"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "process_events 在 empty ring、ring_buff_count=0 时仍进入循环，"
        "触发 assert 或 unsigned underflow"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "ae4dma_unmap_pci_bar unconditionally returns 0 (success) after BAR unmap failure"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "BAR unmap fails; ae4dma_channel_destruct ignores the nonzero rc return value, "
        "frees its retry state, and detach cannot report the failure"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == []
    realloc_result.risks[0].trigger = (
        "vtophys failure leaves a partial descriptor batch with the user callback attached and invoked"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "queue configuration logic is inverted and always selects maximum 16 queues"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    realloc_result.risks[0].trigger = (
        "attach failure has no cleanup path"
    )
    assert _known_c_precheck_order_artifact_ids(
        worker_task,
        realloc_result,
    ) == ["U00-R-SCOPE"]
    review_findings = [
        SimpleNamespace(
            check_id="C01",
            finding="描述符环边界检查在 ring_buff_count 递增之后",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C02",
            finding="ring calloc 分配失败时 qbase_addr 未释放并泄漏",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C03",
            finding="递减 ring_buff_count 前无检查，assert 可能触发",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C05",
            finding="spdk_vtophys 的 size 参数需确认是否正确赋值",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C06",
            finding="qbase_lo/qbase_hi MMIO write 失败后资源泄漏",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C07",
            finding="unmap 硬编码 BAR=0 而非 AE4DMA_PCIE_BAR",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C08",
            finding="ERROR 状态处理后仍推进 tail，导致 descriptor 状态混淆",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C09",
            finding="SUBMITTED 状态 break 会让处理永久停止",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C10",
            finding="spdk_ae4dma_flush 不等待 DMA 完成",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C11",
            finding="attach 失败后没有 cleanup",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C12",
            finding="cmd_q_count 在部分初始化失败时计数不一致",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C13",
            finding="ae4dma_map_pci_bar 返回 -1，不符合 errno 约定",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C14",
            finding="公共 API 返回值 -1、void 与 errno 不一致",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C15",
            finding="硬件描述符错误只记日志不传播，调用方无法感知",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="C16",
            finding=(
                "公共 API 的 ring-full 失败返回正数 1，与 negative errno 契约不一致"
            ),
            evidence=[],
        ),
        SimpleNamespace(
            check_id="CALLBACK-WRONG-DESC",
            finding=(
                "cb_desc 可能不是 last_desc，成功后把 callback 挂到错误 descriptor"
            ),
            evidence=[],
        ),
        SimpleNamespace(
            check_id="QUEUE-INVERTED",
            finding="q_per_eng 的 queue 分支逻辑反置，输入范围与赋值矛盾",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="RING-THRESHOLD",
            finding="ring 的 > 与 >= boundary 阈值不一致",
            evidence=[],
        ),
        SimpleNamespace(
            check_id="VTOPHYS-PARTIAL-CALLBACK",
            finding="vtophys 失败后此前描述符仍带 callback 并触发回调",
            evidence=[],
        ),
    ]
    assert _known_c_review_misread_finding_ids(
        worker_task,
        review_findings,
    ) == [
        "C01", "C02", "C03", "C05", "C06", "C07", "C08",
        "C09", "C10", "C11", "C12", "C13", "C14", "C15",
        "CALLBACK-WRONG-DESC", "QUEUE-INVERTED", "RING-THRESHOLD",
        "VTOPHYS-PARTIAL-CALLBACK",
    ]
    unmap_findings = [
        SimpleNamespace(
            check_id="C04",
            finding="detach 忽略 unmap 返回值并继续释放 channel",
            evidence=[],
        ),
    ]
    assert _known_c_unmap_finding_ids(
        worker_task,
        unmap_findings,
    ) == ["C04"]


def _derived_semantic_path_does_not_rewrite_main_path() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    task["semantic_check_items"] = [{
        "check_id": "SC-LIFECYCLE-01",
        "kind": "runtime_semantics",
        "subject_path": "module/entry.c",
        "instruction": "主路径只检查一次正常生命周期，重试使用派生 path",
        "context_paths": ["module/entry.c"],
    }]
    write_json(task_path, task)
    result = _task_result(task)
    result["risks"] = [{
        "risk_id": "R-RETRY",
        "title": "失败重试后的独立终态",
        "affected_paths": ["module/entry.c"],
        "dfx": ["可靠性与容错"],
        "severity": "Medium",
        "confidence": "medium",
        "trigger": "首次失败后修复条件并重试",
        "system_result": "重试路径保留部分状态",
        "external_observation": "重试后的业务状态可观察",
        "exclusion_condition": "未执行失败重试路径",
        "upstream_semantics": {
            "reachability": "公开入口允许重试",
            "caller_constraints": "调用方未阻断重试",
            "documented_behavior": "资料未定义该残留为预期",
            "existing_tests": "当前无确认用例",
            "conclusion": "unresolved",
        },
        "translation_status": "Developer-confirm",
        "status": "pending",
        "evidence": result["evidence"],
    }]
    result["analysis_checkpoint"]["failure_paths"] = [{
        "path_id": "SC-LIFECYCLE-01",
        "linked_risk_ids": [],
        "trigger": "有效配置下首次正常调用",
        "side_effects": "正常初始化一次",
        "failure": "无故障",
        "caller_handling": "读取成功结果",
        "final_states": "正常可用",
        "disposition": "excluded",
    }, {
        "path_id": "SC-LIFECYCLE-01:retry",
        "linked_risk_ids": ["R-RETRY"],
        "trigger": "首次失败后修复条件并重试",
        "side_effects": "首次副作用残留后新增重试副作用",
        "failure": "重试路径状态未完全恢复",
        "caller_handling": "调用方继续观测状态",
        "final_states": "重试路径保留部分状态",
        "disposition": "risk",
    }]
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(task_path)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"


def _write_all_analysis(state: dict) -> None:
    state = _advance_to_test_generation(state)
    _bind_analysis_actions(state)
    for task_path in state["agent_task_paths"]:
        task = read_json(Path(task_path))
        result = _task_result(task)
        gaps = [
            (context, gap)
            for context in task.get("coverage_context", [])
            for gap in context.get("gaps", [])
        ]
        for index, (context, gap) in enumerate(gaps, start=1):
            case_id = f"{task['unit']['unit_id']}-COV-AUTO-{index:02d}"
            case = _test_case(case_id, title=f"补测零覆盖函数 {context['function']}")
            case["linked_coverage_ids"] = [gap["coverage_id"]]
            result["test_cases"].append(case)
            result["analysis_checkpoint"]["coverage_decisions"].append({
                "coverage_id": gap["coverage_id"],
                "disposition": "new_coverage_case",
                "linked_test_case_ids": [case_id],
                "reason": "通过当前模块业务入口补齐函数级零覆盖路径。",
            })
        write_json(Path(task["result_path"]), result)
        _validate_worker_task(Path(task_path))


def _advance_to_review_comparison(
    data_root: Path,
    run_id: str,
    *,
    reviewer: str = "reviewer-1",
    extra_findings: list[dict] | None = None,
) -> tuple[dict, list[dict]]:
    run_dir = data_root / "runs" / run_id
    task = read_json(run_dir / "agent-tasks" / "review-independent.json")
    assert task["stage"] == "independent_review"
    assert task["analysis_results"] == []
    progress = read_json(run_dir / "progress.json")
    if progress["agent_sessions"]["review"]["task_id"] is None:
        subprocess.run(
            [
                sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
                "--run-id", run_id, "--data-root", str(data_root),
                "--role", "review", "--task-id", _REVIEW_SESSION_ID,
            ],
            check=True,
            capture_output=True,
        )
    findings = []
    for reference in task["analysis_tasks"]:
        worker_task = read_json(Path(reference["task_path"]))
        unit_findings = [{
            "unit_id": reference["unit_id"],
            "check_id": item["check_id"],
            "finding": f"独立核对 {item['check_id']} 未发现额外问题",
            "evidence": [f"{worker_task['unit']['repo_id']}:{item['subject_path']}:1"],
        } for item in worker_task["semantic_check_items"]]
        if not unit_findings:
            unit_findings.append({
                "unit_id": reference["unit_id"],
                "check_id": f"BASELINE-{reference['unit_id']}",
                "finding": "独立核对正常与异常路径，未发现 worker 之外的额外问题",
                "evidence": [
                    f"{worker_task['unit']['repo_id']}:{worker_task['unit']['source_scope'][0]}:1"
                ],
            })
        findings.extend(unit_findings)
    findings.extend(extra_findings or [])
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": run_id,
        "reviewer_id": reviewer,
        "finish_reason": "stop",
        "summary": "独立复核完成",
        "reviewed_units": [item["unit_id"] for item in task["analysis_tasks"]],
        "findings": findings,
    })
    _validate_review_task(run_dir / "agent-tasks" / "review-independent.json")
    state = run_module_analysis(str(run_dir / "inputs" / "task-contract.json"))
    assert state["phase"] == "WAITING_COMPARISON_REVIEW"
    comparison = read_json(run_dir / "agent-tasks" / "review.json")
    assert comparison["stage"] == "comparison_review"
    independent = read_json(run_dir / "agent-results" / "review-independent.json")
    assert comparison["same_reviewer_id"] == independent["reviewer_id"]
    return comparison, findings


def _compared_findings(task: dict, findings: list[dict], disposition_for=None) -> list[dict]:
    payloads = []
    first_for_unit: dict[str, dict] = {}
    for finding in findings:
        disposition = (
            disposition_for(finding) if disposition_for is not None else "covered"
        )
        payload = {
            **finding,
            "worker_disposition": disposition,
            "linked_worker_risk_ids": [],
            "linked_worker_test_case_ids": [],
        }
        payloads.append(payload)
        if disposition != "reasonably_excluded":
            first_for_unit.setdefault(finding["unit_id"], payload)
    for result_ref in task["analysis_results"]:
        target = first_for_unit[result_ref["unit_id"]]
        worker_result = read_json(Path(result_ref["result_path"]))
        target["linked_worker_risk_ids"] = [
            item["risk_id"] for item in worker_result["risks"]
        ]
        target["linked_worker_test_case_ids"] = [
            item["test_case_id"] for item in worker_result["test_cases"]
        ]
    return payloads


def _test_case_checks(task: dict, verdict: str = "valid") -> list[dict]:
    checks: list[dict] = []
    for result_ref in task["analysis_results"]:
        worker_result = read_json(Path(result_ref["result_path"]))
        for case in worker_result["test_cases"]:
            checks.append({
                "unit_id": result_ref["unit_id"],
                "test_case_id": case["test_case_id"],
                "expected_results": [step["expected_result"] for step in case["steps"]],
                "failure_observations": [
                    step.get("failure_observation") for step in case["steps"]
                ],
                "current_behavior": "已按冻结源码独立核对当前行为",
                "verdict": verdict,
                "reason": "步骤预期是正确产品通过标准，不是当前错误实现的复述",
            })
    return checks


def _review(data_root: Path, run_id: str, *, status: str, reviewer: str = "reviewer-1") -> None:
    task, findings = _advance_to_review_comparison(data_root, run_id, reviewer=reviewer)
    issues = [] if status == "PASS" else [{
        "issue_id": "I-001",
        "unit_id": task["analysis_results"][0]["unit_id"],
        "reason": "缺少异常分支说明",
        "required_change": "补充异常分支分析",
    }]
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": run_id,
        "reviewer_id": reviewer,
        "finish_reason": "stop",
        "status": status,
        "summary": "复核完成",
        "issues": issues,
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(task, findings),
        "test_case_checks": _test_case_checks(task),
    })
    _validate_review_task(data_root / "runs" / run_id / "agent-tasks" / "review.json")


def _pass_report() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"
    assert Path(state["report_path"]).is_file() and Path(state["html_report_path"]).is_file()
    markdown = Path(state["report_path"]).read_text(encoding="utf-8")
    html_report = Path(state["html_report_path"]).read_text(encoding="utf-8")
    assert markdown.startswith("# CHAP 分析报告")
    assert "<title>CHAP 分析报告</title>" in html_report
    assert "PANGEA Agent 测试分析报告" not in markdown + html_report
    assert "COMPLETE / PASS" not in markdown + html_report
    assert "| 项目 | 内容 |" in markdown
    assert "| 类别 | 源码仓 | 路径 | 纳入原因 |" in markdown
    assert "## 2. 分析引用范围" in markdown
    assert "质量门禁已通过。完成" in markdown and "质量门禁已通过。完成" in html_report


def _late_worker_rejection_reuses_independent_review() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    run_dir = data_root / "runs" / "smoke-01"
    independent_task_path = run_dir / "agent-tasks" / "review-independent.json"
    independent_task = read_json(independent_task_path)
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "review", "--task-id", _REVIEW_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    findings = []
    for reference in independent_task["analysis_tasks"]:
        worker_task = read_json(Path(reference["task_path"]))
        findings.extend({
            "unit_id": reference["unit_id"],
            "check_id": item["check_id"],
            "finding": f"独立核对 {item['check_id']} 未发现额外问题",
            "evidence": [f"{worker_task['unit']['repo_id']}:{item['subject_path']}:1"],
        } for item in worker_task["semantic_check_items"])
    write_json(Path(independent_task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "summary": "独立复核完成",
        "reviewed_units": [item["unit_id"] for item in independent_task["analysis_tasks"]],
        "findings": findings,
    })
    _validate_review_task(independent_task_path)

    worker_task_path = run_dir / "agent-tasks" / "analysis" / "U00-test_generation.json"
    worker_task = read_json(worker_task_path)
    worker_result_path = Path(worker_task["result_path"])
    worker_result = read_json(worker_result_path)
    worker_result["finish_reason"] = "length"
    write_json(worker_result_path, worker_result)

    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_TEST_GENERATION"
    assert [action["stage"] for action in state["agent_actions"]] == ["test_generation"]
    assert [action["action"] for action in state["agent_actions"]] == ["dispatch_agent"]
    retry_task = read_json(worker_task_path)
    assert retry_task["validation_feedback"]

    worker_result["finish_reason"] = "stop"
    write_json(worker_result_path, worker_result)
    _bind_analysis_actions(state)
    _validate_worker_task(worker_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_COMPARISON_REVIEW"
    assert read_json(run_dir / "agent-tasks" / "review.json")["same_reviewer_id"] == read_json(
        run_dir / "agent-results" / "review-independent.json"
    )["reviewer_id"]


def _review_missing_finding_cannot_pass() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    extra = {
        "unit_id": "U00",
        "check_id": "LIFECYCLE-U00",
        "finding": "独立复核发现一条可达失败路径",
        "evidence": ["demo:module/entry.c:1"],
    }
    task, findings = _advance_to_review_comparison(
        data_root,
        "smoke-01",
        extra_findings=[extra],
    )
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "发现遗漏但错误放行",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(
            task,
            findings,
            lambda finding: (
                "missing" if finding["check_id"] == "LIFECYCLE-U00" else "covered"
            ),
        ),
        "test_case_checks": _test_case_checks(task),
    })
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main",
            "check-review-artifact", "--task",
            str(data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "PASS 不能包含 missing" in rejected.stderr
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert progress["phase"] == "WAITING_COMPARISON_REVIEW"


def _review_blocking_finding_must_reach_rework_issue() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    extra = {
        "unit_id": "U00",
        "check_id": "LIFECYCLE-U00",
        "finding": "独立复核发现一条必须返工的生命周期路径",
        "evidence": ["demo:module/entry.c:1"],
    }
    task, findings = _advance_to_review_comparison(
        data_root,
        "smoke-01",
        extra_findings=[extra],
    )
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "REWORK",
        "summary": "识别到生命周期遗漏",
        "issues": [{
            "issue_id": "I-001",
            "unit_id": "U00",
            "reason": "要求补充遗漏的生命周期路径，但错误地没有标识来源 finding",
            "required_change": "新增对应风险和测试用例",
        }],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(
            task,
            findings,
            lambda finding: (
                "missing" if finding["check_id"] == "LIFECYCLE-U00" else "covered"
            ),
        ),
        "test_case_checks": _test_case_checks(task),
    })
    rejected = _reject_review_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    )
    assert "确保返工任务不会漏项" in rejected


def _comparison_must_account_for_every_worker_artifact() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["risks"] = [{
        "risk_id": "U00-R-ACCOUNT",
        "title": "需要逐项对照的业务风险",
        "affected_paths": [task["unit"]["source_scope"][0]],
        "dfx": ["功能与状态"],
        "severity": "Medium",
        "confidence": "high",
        "trigger": "业务入口触发失败路径",
        "system_result": "模块返回失败",
        "external_observation": "业务返回值可观察",
        "exclusion_condition": "失败条件未触发",
        "upstream_semantics": {
            "reachability": "业务入口可达",
            "caller_constraints": "调用方未消除",
            "documented_behavior": "资料未定义为预期",
            "existing_tests": "已有测试未覆盖",
            "conclusion": "risk_remains",
        },
        "translation_status": "Blackbox-ready",
        "status": "pending",
        "evidence": result["evidence"],
    }]
    result["test_cases"] = [_test_case(
        "U00-TC-ACCOUNT", risk_ids=["U00-R-ACCOUNT"]
    )]
    result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R-ACCOUNT"],
        "failure": "业务操作失败",
        "final_states": "模块返回失败",
        "disposition": "risk",
    })
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    task, findings = _advance_to_review_comparison(data_root, "smoke-01")
    compared = _compared_findings(task, findings)
    for finding in compared:
        finding["linked_worker_risk_ids"] = []
        finding["linked_worker_test_case_ids"] = []
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "错误地跳过 worker 产物逐项对照",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": compared,
        "test_case_checks": _test_case_checks(task),
    })
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main",
            "check-review-artifact", "--task",
            str(data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "linked_worker_test_case_ids" in rejected.stderr
    compared.append({
        "unit_id": "U00",
        "check_id": "WORKER-U00-R-ACCOUNT",
        "finding": "comparison 阶段从 Worker 产物回查源码确认该失败路径成立",
        "evidence": ["demo:module/entry.c:1"],
        "worker_disposition": "covered",
        "linked_worker_risk_ids": ["U00-R-ACCOUNT"],
        "linked_worker_test_case_ids": ["U00-TC-ACCOUNT"],
    })
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "对照补充 finding 承接独立阶段未覆盖的有效 Worker 产物",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": compared,
        "test_case_checks": _test_case_checks(task),
    })
    assert "必须使用 COMPARISON- 前缀" in _reject_review_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    )
    compared[-1]["check_id"] = "COMPARISON-U00-R-ACCOUNT"
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "对照补充 finding 承接独立阶段未覆盖的有效 Worker 产物",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": compared,
        "test_case_checks": _test_case_checks(task),
    })
    _validate_review_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    )
    assert run_module_analysis(str(contract))["phase"] == "COMPLETE"


def _comparison_cannot_invent_coverage_closure() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    task, findings = _advance_to_review_comparison(data_root, "smoke-01")
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "REWORK",
        "summary": "错误要求补充不存在的 Coverage 闭环",
        "issues": [{
            "issue_id": "ISSUE-U00-COVERAGE",
            "unit_id": "U00",
            "reason": "错误声称存在 Coverage gap",
            "required_change": "补充 coverage_decisions 并增加 linked_coverage_ids",
        }],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(task, findings),
        "test_case_checks": _test_case_checks(task),
    })
    review_task_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    assert "不得要求伪造 Coverage 闭环" in _reject_review_task(review_task_path)


def _comparison_cannot_replace_oracle_with_buggy_behavior() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    task, findings = _advance_to_review_comparison(data_root, "smoke-01")
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "REWORK",
        "summary": "错误地把源码缺陷改写成测试通过标准",
        "issues": [{
            "issue_id": "ISSUE-U00-BUGGY-ORACLE",
            "unit_id": "U00",
            "reason": "把需求正确值误判为用例错误",
            "required_change": (
                "修正 TC-U00-01 的 expected_result 为值 2；当前实现违反 REQ-DEMO-001"
            ),
        }],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(task, findings),
        "test_case_checks": _test_case_checks(task),
    })
    review_task_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    assert "不得把当前错误实现写成测试通过标准" in _reject_review_task(review_task_path)

    payload = read_json(Path(task["result_path"]))
    payload["issues"] = [{
        "issue_id": "ISSUE-U00-CONFLATED-ORACLE",
        "unit_id": "U00",
        "reason": (
            "expected_result 与 failure_observation 对同一字段互相否定，"
            "错误地把正确预期与当前失败观测的差异当成用例矛盾"
        ),
        "required_change": "按正确契约修正 expected_result",
    }]
    write_json(Path(task["result_path"]), payload)
    assert (
        "不得要求 expected_result 与 failure_observation 相同或不矛盾"
        in _reject_review_task(review_task_path)
    )
    payload["issues"] = [{
        "issue_id": "ISSUE-U00-ASAN-ORACLE",
        "unit_id": "U00",
        "reason": "错误地把缺陷信号当作故障注入步骤的通过标准",
        "required_change": "将 expected_result 改为：ASan 报告 double-free 或进程崩溃",
    }]
    write_json(Path(task["result_path"]), payload)
    assert "不得把当前错误实现写成测试通过标准" in _reject_review_task(
        review_task_path
    )
    payload["issues"] = [{
        "issue_id": "ISSUE-U00-ACTUAL-BUGGY-ORACLE",
        "unit_id": "U00",
        "reason": "错误地把当前缺陷结果要求成通过标准",
        "required_change": (
            "Rewrite TC-U00-01 expected_result to describe the actual buggy behavior"
        ),
    }]
    write_json(Path(task["result_path"]), payload)
    assert "不得把当前错误实现写成测试通过标准" in _reject_review_task(
        review_task_path
    )


def _comparison_does_not_rework_worker_for_reviewer_mistake() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    task, findings = _advance_to_review_comparison(data_root, "smoke-01")
    base = {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(task, findings),
        "test_case_checks": _test_case_checks(task),
    }
    write_json(Path(task["result_path"]), {
        **base,
        "status": "REWORK",
        "summary": "错误地让 worker 修正 reviewer 自己的误判",
        "issues": [{
            "issue_id": "ISSUE-U00-REVIEWER-MISTAKE",
            "unit_id": "U00",
            "reason": "独立 finding 读错冻结源码，Worker 结论与源码一致",
            "required_change": "修正 SC-DEMO-01 的 finding，改成与冻结源码一致",
        }],
    })
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main",
            "check-review-artifact", "--task",
            str(data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "不得派发 worker 返工" in rejected.stderr

    write_json(Path(task["result_path"]), {
        **base,
        "status": "PASS",
        "summary": "独立 finding 已被冻结源码推翻，worker 结论正确",
        "issues": [],
    })
    _validate_review_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    )
    assert run_module_analysis(str(contract))["phase"] == "COMPLETE"


def _comparison_cannot_drop_independent_findings() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    extra = {
        "unit_id": "U00",
        "check_id": "LIFECYCLE-U00",
        "finding": "独立复核记录正常生命周期",
        "evidence": ["demo:module/entry.c:1"],
    }
    task, _ = _advance_to_review_comparison(
        data_root,
        "smoke-01",
        extra_findings=[extra],
    )
    write_json(Path(task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "错误丢弃独立 finding",
        "issues": [],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": _test_case_checks(task),
    })
    review_task_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    assert "必须逐项保留独立复核 findings" in _reject_review_task(review_task_path)


def _rework_same_reviewer() -> None:
    _, data_root, contract = _lua_workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="REWORK", reviewer="reviewer-1")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK"
    rework_task = read_json(data_root / "runs" / "smoke-01" / "agent-tasks" / "rework" / "U00.json")
    analysis_task = read_json(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "analysis"
        / "U00-test_generation.json"
    )
    assert rework_task["checkpoint_rubric_paths"] == analysis_task["checkpoint_rubric_paths"]
    assert "src/pangea_agent/rubrics/builtin/lua_analysis.md" in rework_task["checkpoint_rubric_paths"]
    assert "src/pangea_agent/rubrics/builtin/openubmc_analysis.md" in rework_task["checkpoint_rubric_paths"]
    prior_path = Path(rework_task["prior_result_path"])
    prior_text = prior_path.read_text(encoding="utf-8")
    prior_result = read_json(prior_path)
    assert Path(rework_task["result_path"]).exists()
    graph_prepared_rework = read_json(Path(rework_task["result_path"]))
    for field in ("summary", "evidence", "business_flows", "risks", "test_cases", "analysis_checkpoint"):
        assert graph_prepared_rework[field] == prior_result[field]
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "prepare-worker-result",
            "--task", str(data_root / "runs" / "smoke-01" / "agent-tasks" / "rework" / "U00.json"),
        ],
        check=True,
        capture_output=True,
    )
    assert prior_path.read_text(encoding="utf-8") == prior_text
    prepared_rework = read_json(Path(rework_task["result_path"]))
    for field in ("summary", "evidence", "business_flows", "risks", "test_cases", "analysis_checkpoint"):
        assert prepared_rework[field] == prior_result[field]
    assert prepared_rework["attempt"] == 1
    assert prepared_rework["completed_stage"] == "rework"
    assert prepared_rework["addressed_review_issue_ids"] == []
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "rework", "--unit-id", "U00", "--task-id", _REPLACEMENT_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "rework", "--unit-id", "U00", "--task-id", _SECOND_REPLACEMENT_SESSION_ID,
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0 and "不能替换 task_id" in rejected.stderr
    write_json(Path(rework_task["result_path"]), _task_result(rework_task))
    _validate_worker_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "rework" / "U00.json"
    )
    prior_review_path = (
        data_root / "runs" / "smoke-01" / "agent-results" / "review.json"
    )
    prior_review = read_json(prior_review_path)
    prior_review["independent_findings"][0]["linked_worker_risk_ids"] = ["REMOVED-RISK"]
    prior_review["independent_findings"][0]["linked_worker_test_case_ids"] = ["REMOVED-TEST"]
    write_json(prior_review_path, prior_review)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK_VERIFICATION"
    review_task = read_json(data_root / "runs" / "smoke-01" / "agent-tasks" / "rework-review.json")
    assert review_task["same_reviewer_id"] == _REVIEW_SESSION_ID
    assert review_task["independent_result_path"].endswith("review-independent.json")
    graph_prepared_review = read_json(Path(review_task["result_path"]))
    assert all(
        "REMOVED-RISK" not in item["linked_worker_risk_ids"]
        and "REMOVED-TEST" not in item["linked_worker_test_case_ids"]
        for item in graph_prepared_review["independent_findings"]
    )
    write_json(Path(review_task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": _REVIEW_SESSION_ID,
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "返工验证通过",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in review_task["analysis_results"]],
        "independent_findings": [],
    })
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main",
            "check-review-artifact", "--task",
            str(data_root / "runs" / "smoke-01" / "agent-tasks" / "rework-review.json"),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "逐项保留独立复核 findings" in rejected.stderr
    independent = read_json(Path(review_task["independent_result_path"]))
    write_json(Path(review_task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": _REVIEW_SESSION_ID,
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "返工验证通过",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in review_task["analysis_results"]],
        "independent_findings": _compared_findings(
            review_task, independent["findings"]
        ),
        "test_case_checks": _test_case_checks(review_task),
    })
    _validate_review_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "rework-review.json"
    )
    assert run_module_analysis(str(contract))["phase"] == "COMPLETE"


def _rework_preserves_allowed_material_paths() -> None:
    _, data_root, contract = _workspace()
    inbox = data_root / "inbox"
    inbox.mkdir(parents=True, exist_ok=True)
    material_path = "inbox/rework-requirement.md"
    (data_root / material_path).write_text(
        "# Rework requirement\n\nThe component must recover after initialization failure.\n",
        encoding="utf-8",
    )
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    _review(data_root, "smoke-01", status="REWORK")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK"
    run_dir = data_root / "runs" / "smoke-01"
    analysis_task = read_json(
        run_dir / "agent-tasks" / "analysis" / "U00-test_generation.json"
    )
    rework_task_path = run_dir / "agent-tasks" / "rework" / "U00.json"
    rework_task = read_json(rework_task_path)
    assert material_path in analysis_task["allowed_material_paths"]
    assert rework_task["allowed_material_paths"] == analysis_task["allowed_material_paths"]
    rework_task["allowed_material_paths"] = []
    write_json(rework_task_path, rework_task)
    repaired_task = load_worker_task(rework_task_path)
    assert repaired_task.allowed_material_paths == analysis_task["allowed_material_paths"]
    rework_task = repaired_task.model_dump(mode="json")
    write_json(Path(rework_task["result_path"]), _task_result(rework_task))
    _validate_worker_task(rework_task_path)


def _invalid_rework_finishes_unresolved_without_reviewer_loop() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    _review(data_root, "smoke-01", status="REWORK")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK"
    run_dir = data_root / "runs" / "smoke-01"
    rework_task_path = run_dir / "agent-tasks" / "rework" / "U00.json"
    rework_task = read_json(rework_task_path)
    write_json(Path(rework_task["result_path"]), _task_result(rework_task))
    _validate_worker_task(rework_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK_VERIFICATION"
    review_task_path = run_dir / "agent-tasks" / "rework-review.json"
    review_task = read_json(review_task_path)
    independent = read_json(Path(review_task["independent_result_path"]))
    write_json(Path(review_task["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": _REVIEW_SESSION_ID,
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "返工验证通过",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in review_task["analysis_results"]],
        "independent_findings": _compared_findings(
            review_task, independent["findings"]
        ),
        "test_case_checks": _test_case_checks(review_task),
    })
    _validate_review_task(review_task_path)
    invalid_rework = read_json(Path(rework_task["result_path"]))
    invalid_rework["summary"] = ""
    write_json(Path(rework_task["result_path"]), invalid_rework)
    state = run_module_analysis(str(contract))
    if state["phase"] == "READY_TO_FINALIZE":
        state = run_module_analysis(str(contract))
    assert state["phase"] == "INCOMPLETE"
    assert state["quality_report"]["status"] == "UNRESOLVED"
    progress = read_json(run_dir / "progress.json")
    assert any(error["kind"] == "rework_result_rejected" for error in progress["errors"])


def _truncated_correction() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    task_path = Path(state["agent_task_paths"][0])
    write_json(Path(task["result_path"]), _task_result(task, finish_reason="truncated"))
    assert "finish_reason=truncated" in _reject_worker_task(task_path)
    assert read_json(data_root / "runs" / "smoke-01" / "progress.json")["phase"] == "WAITING_TEST_GENERATION"
    write_json(Path(task["result_path"]), _task_result(task))
    _validate_worker_task(task_path)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"


def _reviewer_unavailable_has_explicit_unresolved_path() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    _review(data_root, "smoke-01", status="REWORK", reviewer="reviewer-1")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK"
    rework_task = read_json(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "rework" / "U00.json"
    )
    write_json(Path(rework_task["result_path"]), _task_result(rework_task))
    _validate_worker_task(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "rework" / "U00.json"
    )
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK_VERIFICATION"

    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "mark-reviewer-unavailable",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--reviewer-id", "another-reviewer", "--reason", "session missing",
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0 and "不是当前 Run 绑定的原 reviewer" in rejected.stderr
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "mark-reviewer-unavailable",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--reviewer-id", _REVIEW_SESSION_ID,
            "--reason", "saved DSH session cannot be resumed",
        ],
        check=True,
        capture_output=True,
    )
    state = run_module_analysis(str(contract))
    if state["phase"] == "READY_TO_FINALIZE":
        state = run_module_analysis(str(contract))
    assert state["quality_report"]["status"] == "UNRESOLVED"
    assert state["phase"] == "INCOMPLETE"


def _unmatched_evidence_is_pending() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    forged = _task_result(task, fake_location=True)
    forged["evidence"][0]["chunk_id"] = "missing-chunk"
    forged["business_flows"][0]["evidence"] = forged["evidence"]
    write_json(Path(task["result_path"]), forged)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    rejected = read_json(Path(task["result_path"]))
    assert rejected["evidence"][0]["chunk_id"] == "missing-chunk"
    assert "status" not in rejected["evidence"][0]
    assert rejected["business_flows"][0]["evidence"][0]["chunk_id"] == "missing-chunk"
    assert "证据未绑定当前 Run 的真实索引片段" in rejection


def _pass_rejects_unconfirmed_risk_evidence() -> None:
    _, data_root, contract = _workspace()
    state = _advance_to_test_generation(run_module_analysis(str(contract)))
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    pending_evidence = {
        **result["evidence"][0],
        "chunk_id": "missing-risk-chunk",
        "location": None,
    }
    result["risks"] = [{
        "risk_id": "U00-R-PENDING",
        "title": "证据未确认风险",
        "affected_paths": task["unit"]["source_scope"],
        "dfx": ["功能与状态"],
        "severity": "Medium",
        "confidence": "low",
        "trigger": "执行目标业务操作",
        "system_result": "系统返回错误状态",
        "external_observation": "调用方观察到失败",
        "exclusion_condition": "未执行目标业务操作",
        "upstream_semantics": {
            "reachability": "业务入口可达",
            "caller_constraints": "调用方未阻断",
            "documented_behavior": "资料要求操作成功",
            "existing_tests": "未验证该路径",
            "conclusion": "risk_remains",
        },
        "translation_status": "Blackbox-ready",
        "status": "pending",
        "evidence": [pending_evidence],
    }]
    result["test_cases"] = [_test_case(
        "U00-TC-PENDING", risk_ids=["U00-R-PENDING"]
    )]
    result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R-PENDING"],
        "failure": "系统返回错误状态",
        "final_states": "调用方观察到失败",
        "disposition": "risk",
    })
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    rejected = read_json(Path(task["result_path"]))
    assert rejected["risks"][0]["evidence"][0]["chunk_id"] == "missing-risk-chunk"
    assert "status" not in rejected["risks"][0]["evidence"][0]
    assert "证据未绑定当前 Run 的真实索引片段" in rejection


def _comparison_requires_test_case_oracle_checks() -> None:
    _, data_root, contract = _workspace()
    state = _advance_to_test_generation(run_module_analysis(str(contract)))
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["test_cases"] = [_test_case(
        "U00-TC-ORACLE", requirement_ids=["REQ-ORACLE-01"]
    )]
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    review_task, findings = _advance_to_review_comparison(data_root, "smoke-01")
    skeleton = review_result_skeleton(
        ReviewTask.model_validate(review_task),
        IndependentReviewResult.model_validate(
            read_json(Path(review_task["independent_result_path"]))
        ),
    )
    assert skeleton["test_case_checks"] == [{
        "unit_id": "U00",
        "test_case_id": "U00-TC-ORACLE",
        "expected_results": [
            step["expected_result"] for step in result["test_cases"][0]["steps"]
        ],
        "failure_observations": [
            step.get("failure_observation") for step in result["test_cases"][0]["steps"]
        ],
        "current_behavior": "待 reviewer 按冻结源码独立填写",
        "verdict": "unresolved",
        "reason": "待 reviewer 对照正确产品通过标准与当前实现行为",
    }]
    base = {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "复核通过标准",
        "issues": [],
        "reviewed_units": ["U00"],
        "independent_findings": _compared_findings(review_task, findings),
    }
    review_task_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    write_json(Path(review_task["result_path"]), {**base, "test_case_checks": []})
    assert "逐条检查全部 TestCase" in _reject_review_task(review_task_path)
    checks = _test_case_checks(review_task, verdict="invalid")
    write_json(Path(review_task["result_path"]), {**base, "test_case_checks": checks})
    assert "PASS 不能包含 invalid" in _reject_review_task(review_task_path)
    write_json(Path(review_task["result_path"]), {
        **base,
        "status": "REWORK",
        "issues": [{
            "issue_id": "I-ORACLE-UNASSIGNED",
            "unit_id": "U00",
            "reason": "有一条测试通过标准与源码不符",
            "required_change": "修正对应测试的通过标准",
        }],
        "test_case_checks": checks,
    })
    assert "issue 明确点名" in _reject_review_task(review_task_path)
    checks[0]["verdict"] = "valid"
    missing_links = _compared_findings(review_task, findings)
    missing_links[0]["linked_worker_test_case_ids"] = []
    write_json(Path(review_task["result_path"]), {
        **base,
        "independent_findings": missing_links,
        "test_case_checks": checks,
    })
    assert "linked_worker_test_case_ids" in _reject_review_task(review_task_path)
    write_json(Path(review_task["result_path"]), {**base, "test_case_checks": checks})
    _validate_review_task(review_task_path)
    assert run_module_analysis(str(contract))["phase"] == "COMPLETE"


def _duplicate_ids_correction() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    tasks = [read_json(Path(path)) for path in state["agent_task_paths"]]
    for task in tasks:
        result = _task_result(task)
        evidence = result["evidence"]
        result["risks"] = [{
            "risk_id": "DUP-R01",
            "title": "重复风险编号",
            "affected_paths": task["unit"]["source_scope"],
            "dfx": ["功能与状态"],
            "severity": "Medium",
            "confidence": "high",
            "trigger": "触发条件",
            "system_result": "系统结果",
            "external_observation": "外部观测",
            "exclusion_condition": "排除条件",
            "upstream_semantics": {
                "reachability": "业务入口可达",
                "caller_constraints": "调用方没有消除该结果",
                "documented_behavior": "资料未定义为预期行为",
                "existing_tests": "已有测试未覆盖该条件",
                "conclusion": "risk_remains",
            },
            "translation_status": "Blackbox-ready",
            "status": "pending",
            "evidence": evidence,
        }]
        result["analysis_checkpoint"]["failure_paths"][0].update({
            "linked_risk_ids": ["DUP-R01"],
            "failure": "当前单元业务失败",
            "final_states": "当前单元状态异常",
            "disposition": "risk",
        })
        result["test_cases"] = [_test_case("DUP-TC01", risk_ids=["DUP-R01"])]
        write_json(Path(task["result_path"]), result)
    for task_path in state["agent_task_paths"]:
        _validate_worker_task(Path(task_path))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    normalized = [read_json(Path(task["result_path"])) for task in tasks]
    assert {item["risks"][0]["risk_id"] for item in normalized} == {"DUP-R01", "DUP-R01-2"}
    assert {item["test_cases"][0]["test_case_id"] for item in normalized} == {"DUP-TC01", "DUP-TC01-2"}
    assert not read_json(data_root / "runs" / "smoke-01" / "progress.json")["errors"]


def _mechanical_task_change_does_not_block() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    task["result_path"] = "Z:\\wrong\\result.json"
    write_json(task_path, task)
    corrected_path = Path(task_path).parents[2] / "agent-results" / "analysis" / "U00.json"
    write_json(corrected_path, _task_result(task))
    _validate_worker_task(task_path)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    assert Path(read_json(task_path)["result_path"]).resolve() == corrected_path.resolve()


def _mechanical_review_fields_do_not_block() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"
    run_dir = data_root / "runs" / "smoke-01"
    _, findings = _advance_to_review_comparison(data_root, "smoke-01")
    task_path = run_dir / "agent-tasks" / "review.json"
    task = read_json(task_path)
    task["result_path"] = "Z:\\wrong\\review.json"
    write_json(task_path, task)
    write_json(run_dir / "agent-results" / "review.json", {
        "schema_version": "1.0",
        "run_id": "wrong-run",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "语义复核通过",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in task["analysis_results"]],
        "independent_findings": _compared_findings(task, findings),
        "test_case_checks": _test_case_checks(task),
    })
    _validate_review_task(task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"
    normalized = read_json(run_dir / "agent-results" / "review.json")
    assert normalized["run_id"] == "smoke-01"


def _missing_scope() -> None:
    _, data_root, contract = _workspace()
    payload = read_json(contract)
    payload["source_scope"] = ["missing"]
    write_json(contract, payload)
    try:
        run_module_analysis(str(contract))
    except ValueError as exc:
        assert "没有可分析源码" in str(exc)
    else:
        raise AssertionError("不存在的源码范围仍派发了任务")
    assert not list((data_root / "runs" / "smoke-01" / "agent-tasks").glob("**/*.json"))


def _report_recovery() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"
    Path(state["report_path"]).unlink()
    Path(state["html_report_path"]).unlink()
    recovered = run_module_analysis(str(contract))
    assert recovered["phase"] == "COMPLETE"
    assert Path(recovered["report_path"]).is_file() and Path(recovered["html_report_path"]).is_file()
    run_dir = data_root / "runs" / "smoke-01"
    (run_dir / "final-state.json").unlink()
    Path(recovered["report_path"]).unlink()
    Path(recovered["html_report_path"]).unlink()
    recovered = run_module_analysis(str(contract))
    assert recovered["phase"] == "COMPLETE"
    assert (run_dir / "final-state.json").is_file()
    assert Path(recovered["report_path"]).is_file() and Path(recovered["html_report_path"]).is_file()


def _document_gap() -> None:
    root, data_root, contract = _workspace()
    inbox = data_root / "inbox"
    inbox.mkdir()
    (inbox / "broken.pdf").write_bytes(b"not a pdf")
    image_path = root / "diagram.png"
    Image.new("RGB", (2, 2), "white").save(image_path)
    document = Document()
    document.add_paragraph("CHAP 设计图")
    document.add_picture(str(image_path))
    document.save(inbox / "design.docx")
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "INCOMPLETE"
    assert state["quality_report"]["status"] == "UNRESOLVED"
    kinds = {item.get("kind") for item in state["quality_report"]["unresolved"]}
    assert {"document_parse_warning", "unread_image"} <= kinds


def _coverage_reference_only() -> None:
    _, data_root, contract = _workspace()
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["模块", "函数", "覆盖次数"])
    sheet.append(["module", "demo_start", 3])
    sheet.append(["module", "not_in_scope", 1])
    workbook.save(coverage / "coverage.xlsx")
    state = run_module_analysis(str(contract))
    report = state["coverage_report"]
    assert report["matched"] == [] and report["unmatched"] == []
    source_task = read_json(Path(state["agent_task_paths"][0]))
    assert source_task["coverage_context"] == []
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    assert task["coverage_context"] == []


def _coverage_gap_requires_closure_and_linkage_is_strict() -> None:
    _, data_root, contract = _workspace()
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["模块", "函数", "覆盖次数"])
    sheet.append(["module", "demo_start", 0])
    workbook.save(coverage / "coverage.xlsx")
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    gap_id = task["coverage_context"][0]["gaps"][0]["coverage_id"]
    result = _task_result(task)
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "Coverage 缺口尚未逐项闭环" in rejection
    assert gap_id in rejection

    _, data_root, contract = _workspace()
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["模块", "函数", "覆盖次数"])
    sheet.append(["module", "demo_start", 0])
    workbook.save(coverage / "coverage.xlsx")
    state = _advance_to_test_generation(run_module_analysis(str(contract)))
    task = read_json(Path(state["agent_task_paths"][0]))
    gap_id = task["coverage_context"][0]["gaps"][0]["coverage_id"]
    result = _task_result(task)
    result["test_cases"] = [
        _test_case("U00-COV-TC1", requirement_ids=["REQ-COV-01"], title="Coverage 缺口补测一"),
        _test_case("U00-COV-TC2", requirement_ids=["REQ-COV-02"], title="Coverage 缺口补测二"),
    ]
    result["analysis_checkpoint"]["coverage_decisions"] = [{
        "coverage_id": gap_id,
        "disposition": "new_coverage_case",
        "linked_test_case_ids": ["U00-COV-TC1", "U00-COV-TC2"],
        "reason": "通过当前模块业务入口补齐未执行函数路径。",
    }]
    write_json(Path(task["result_path"]), result)
    rejection = _reject_worker_task(Path(state["agent_task_paths"][0]))
    assert "Coverage 双向闭环存在问题" in rejection
    assert "U00-COV-TC1" in rejection and "U00-COV-TC2" in rejection
    for case in result["test_cases"]:
        case["linked_coverage_ids"] = [gap_id]
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"


def _coverage_report_only_shows_current_zero_functions() -> None:
    _, data_root, contract = _workspace()
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["模块", "函数", "覆盖次数"])
    sheet.append(["module", "demo_start", 0])
    sheet.append(["other", "not_in_scope", 0])
    sheet.append(["module", "already_executed", 3])
    workbook.save(coverage / "coverage.xlsx")
    state = run_module_analysis(str(contract))
    assert [item["function"] for item in state["coverage_report"]["matched"]] == ["demo_start"]
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="PASS")
    complete = run_module_analysis(str(contract))
    assert complete["phase"] == "COMPLETE"
    for report_path in (complete["report_path"], complete["html_report_path"]):
        report = Path(report_path).read_text(encoding="utf-8")
        assert "demo_start" in report
        assert "not_in_scope" not in report
        assert "already_executed" not in report


def _empty_coverage_rejects_claimed_gap() -> None:
    _, data_root, contract = _workspace()

    def unsupported_risk(task: dict, result: dict) -> dict:
        return {
            "risk_id": "U00-R-NO-COVERAGE",
            "title": "无 Coverage 输入时的候选风险",
            "affected_paths": task["unit"]["source_scope"],
            "dfx": ["功能与状态"],
            "severity": "Low",
            "confidence": "low",
            "trigger": "业务入口调用",
            "system_result": "需要开发确认",
            "external_observation": "当前输入无法确认",
            "exclusion_condition": "补充运行证据后排除",
            "upstream_semantics": {
                "reachability": "业务入口可达",
                "caller_constraints": "调用方未消除",
                "documented_behavior": "资料未定义",
                "existing_tests": "Component:init函数的失败路径未在上游测试中触发",
                "conclusion": "unresolved",
            },
            "translation_status": "Developer-confirm",
            "status": "pending",
            "evidence": result["evidence"],
        }

    state = run_module_analysis(str(contract))
    _bind_analysis_actions(state)
    source_task_path = Path(state["agent_task_paths"][0])
    source_task = read_json(source_task_path)
    write_json(Path(source_task["result_path"]), _task_result(source_task))
    _validate_worker_task(source_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"

    risk_task = read_json(Path(state["agent_task_paths"][0]))
    assert risk_task["coverage_context"] == []
    risk_result = _task_result(risk_task)
    risk_result["risks"] = [unsupported_risk(risk_task, risk_result)]
    risk_result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R-NO-COVERAGE"],
        "failure": "当前输入无法确认故障终态",
        "final_states": "需要开发确认",
        "disposition": "unresolved",
    })
    risk_result["analysis_checkpoint"]["coverage_priorities"] = [
        "Component:init 函数覆盖率 0"
    ]
    write_json(Path(risk_task["result_path"]), risk_result)
    assert "coverage_context 为空" in _reject_worker_task(
        Path(state["agent_task_paths"][0])
    )

    risk_result["analysis_checkpoint"]["coverage_priorities"] = []
    risk_result["risks"][0]["upstream_semantics"]["existing_tests"] = "当前单元未提供已有测试证据"
    write_json(Path(risk_task["result_path"]), risk_result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_TEST_GENERATION"

    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["risks"] = [unsupported_risk(task, result)]
    result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R-NO-COVERAGE"],
        "failure": "当前输入无法确认故障终态",
        "final_states": "需要开发确认",
        "disposition": "unresolved",
    })
    write_json(Path(task["result_path"]), result)
    assert "coverage_context 为空" in _reject_worker_task(
        Path(state["agent_task_paths"][0])
    )

    result["risks"][0]["upstream_semantics"]["existing_tests"] = "当前单元未提供已有测试证据"
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    assert run_module_analysis(str(contract))["phase"] == "WAITING_INDEPENDENT_REVIEW"


def _risk_stage_can_plan_coverage_before_test_ids_exist() -> None:
    _, data_root, contract = _workspace()
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["模块", "函数", "覆盖次数"])
    sheet.append(["module", "demo_start", 0])
    workbook.save(coverage / "coverage.xlsx")
    workbook.close()

    state = run_module_analysis(str(contract))
    _bind_analysis_actions(state)
    source_task_path = Path(state["agent_task_paths"][0])
    source_task = read_json(source_task_path)
    write_json(Path(source_task["result_path"]), _task_result(source_task))
    _validate_worker_task(source_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"

    task_path = Path(state["agent_task_paths"][0])
    risk_task = read_json(task_path)
    gap_id = risk_task["coverage_context"][0]["gaps"][0]["coverage_id"]
    risk_result = _task_result(risk_task)
    risk_result["analysis_checkpoint"]["coverage_decisions"] = [{
        "coverage_id": gap_id,
        "disposition": "new_coverage_case",
        "linked_test_case_ids": [],
        "reason": "风险阶段只确定补测方向，测试阶段再生成并关联真实用例。",
    }]
    write_json(Path(risk_task["result_path"]), risk_result)
    subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result", "--task", str(task_path)],
        check=True,
        capture_output=True,
    )
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_TEST_GENERATION"

    test_task_path = Path(state["agent_task_paths"][0])
    assert test_task_path != task_path
    test_task = read_json(test_task_path)
    test_result = _task_result(test_task)
    test_result["analysis_checkpoint"]["coverage_decisions"] = risk_result["analysis_checkpoint"]["coverage_decisions"]
    write_json(Path(test_task["result_path"]), test_result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(test_task_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "未关联测试用例" in rejected.stderr


def _material_traceability_report() -> None:
    _, data_root, contract = _workspace()
    inbox = data_root / "inbox"
    inbox.mkdir()
    (inbox / "current.md").write_text("REQ-DEMO-01 当前需求。\n", encoding="utf-8")
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    assert task["allowed_material_paths"] == ["inbox/current.md"]
    result = _task_result(task)
    result["analysis_checkpoint"]["material_decisions"] = [{
        "path": "Component.changed 类表共享",
        "decision": "current",
        "reason": "不能用源码描述冒充资料清单路径。",
    }]
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(Path(state["agent_task_paths"][0])),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "资料 decision 引用了当前 Run 清单外的路径" in rejected.stderr
    result["analysis_checkpoint"]["material_decisions"] = [{
        "path": "inbox/current.md",
        "decision": "current",
        "reason": "当前需求基线。",
    }]
    result["test_cases"] = [_test_case(
        "U00-MAT-TC1",
        requirement_ids=["REQ-DEMO-01"],
        material_ids=["MAT:inbox/current.md"],
        title="当前需求行为验证",
    )]
    for evidence in result["evidence"] + result["business_flows"][0]["evidence"]:
        evidence["chunk_id"] = "E-RISK-FAKE"
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(Path(state["agent_task_paths"][0])),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "证据未绑定当前 Run 的真实索引片段" in rejected.stderr
    source_chunk_id = f"{task['unit']['repo_id']}:{task['unit']['source_scope'][0]}:1-1"
    for evidence in result["evidence"] + result["business_flows"][0]["evidence"]:
        evidence["chunk_id"] = source_chunk_id
    result["evidence"].append({
        "chunk_id": "material:inbox/current.md:1-1",
        "location": "inbox/current.md:1-1",
        "observation": "REQ-DEMO-01 作为当前需求采用。",
    })
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    review_task_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review-independent.json"
    material = subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "read-material",
            "--task",
            str(review_task_path),
            "--path",
            "inbox/current.md",
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    assert "REQ-DEMO-01 当前需求" in material.stdout
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    markdown = Path(state["report_path"]).read_text(encoding="utf-8")
    html_report = Path(state["html_report_path"]).read_text(encoding="utf-8")
    for expected in ("inbox/current.md", "inbox/current.md:1-1", "REQ-DEMO-01 作为当前需求采用。", "MAT:inbox/current.md"):
        assert expected in markdown and expected in html_report
    assert "### 资料采用与排除结论" not in markdown
    assert "### 资料引用" in markdown


def _confirmed_historical_issues_are_frozen_as_references() -> None:
    _, data_root, contract = _workspace()
    inbox = data_root / "inbox"
    inbox.mkdir()
    (inbox / "current.md").write_text(
        "REQ-HIST-01 重连成功后，新实例不得接收旧会话状态。\n",
        encoding="utf-8",
    )
    catalog_root = data_root / "asset-catalog"
    (catalog_root / "historical-issues").mkdir(parents=True)
    (catalog_root / "historical-issues" / "draft-asset.json").write_text(
        '{"issues":[{"issue_id":"draft-only-issue"}]}\n',
        encoding="utf-8",
    )
    confirmed_id = "asset-incident-r1-issue-001"
    excluded_id = "asset-incident-r1-issue-002"
    write_json(catalog_root / "historical-issue-reviews.json", {
        "schema_version": "1.0",
        "updated_at": "2026-08-24T00:00:00Z",
        "reviews": {
            confirmed_id: {
                "issue_id": confirmed_id,
                "asset_id": "asset-incident",
                "decision": "confirmed",
                "reviewed_at": "2026-08-24T00:00:00Z",
                "issue": {
                    "issue_id": confirmed_id,
                    "status": "confirmed",
                    "non_binding": True,
                    "review_required": False,
                    "title": "重连清理残留",
                    "symptom": "重连后旧状态影响新实例。",
                    "trigger_conditions": ["首次连接清理失败后再次连接"],
                    "impact": ["新实例收到旧状态"],
                    "root_causes": ["清理完成前允许重连"],
                    "resolutions": ["等待清理完成后再重连"],
                    "verification": ["新实例不接收旧状态"],
                    "limitations": [],
                    "missing_fields": [],
                    "confidence": "high",
                    "evidence": [{
                        "location": "inbox/incident.md#line=3",
                        "excerpt": "cleanup timed out",
                    }],
                },
            },
            excluded_id: {
                "issue_id": excluded_id,
                "asset_id": "asset-incident",
                "decision": "excluded",
                "reviewed_at": "2026-08-24T00:00:00Z",
            },
        },
    })
    write_json(catalog_root / "methodology-candidates.json", {
        "schema_version": "1.0",
        "candidates": [{"title": "本轮不应进入 Run"}],
    })

    state = run_module_analysis(str(contract))
    run_dir = data_root / "runs" / "smoke-01"
    historical_path = f"historical-issues/{confirmed_id}.md"
    frozen_history = run_dir / "inputs" / "frozen" / "materials" / historical_path
    assert frozen_history.is_file()
    frozen_text = frozen_history.read_text(encoding="utf-8")
    assert "重连清理残留" in frozen_text
    assert "历史问题参考" in frozen_text
    assert not (frozen_history.parent / f"{excluded_id}.md").exists()
    assert not (frozen_history.parent / "draft-only-issue.md").exists()
    assert not (run_dir / "inputs" / "frozen" / "materials" / "methodology-candidates.json").exists()
    manifest = read_json(run_dir / "inputs" / "source-manifest.json")
    catalog_paths = {item["path"] for item in manifest["material_catalog"]}
    assert historical_path in catalog_paths
    assert "historical-issues/draft-only-issue.md" not in catalog_paths

    _bind_analysis_actions(state)
    source_task_path = Path(state["agent_task_paths"][0])
    source_task = read_json(source_task_path)
    write_json(Path(source_task["result_path"]), _task_result(source_task))
    _validate_worker_task(source_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"

    risk_task_path = Path(state["agent_task_paths"][0])
    risk_task = read_json(risk_task_path)
    material = subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "read-material",
            "--task",
            str(risk_task_path),
            "--path",
            historical_path,
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    chunks = json.loads(material.stdout)
    assert chunks
    assert "historical_issue" in chunks[0]["tags"]
    assert chunks[0]["evidence_role"] == "reference_only"
    assert "重连清理残留" in chunks[0]["content"]

    risk_result = _task_result(risk_task)
    risk_result["analysis_checkpoint"]["material_decisions"] = [{
        "path": "inbox/current.md",
        "decision": "current",
        "reason": "当前需求给出重连通过标准。",
    }, {
        "path": historical_path,
        "decision": "current",
        "reason": "人工确认的历史问题与当前重连流程相关。",
    }]
    historical_evidence = {
        "chunk_id": chunks[0]["chunk_id"],
        "location": chunks[0]["location"],
        "observation": "已确认历史问题“重连清理残留”用于重放当前触发。",
    }
    risk_result["evidence"].append(historical_evidence)
    risk_result["risks"] = [{
        "risk_id": "U00-R-HISTORY",
        "title": "重连前清理未完成",
        "affected_paths": risk_task["unit"]["source_scope"],
        "dfx": ["可靠性与容错"],
        "severity": "Medium",
        "confidence": "medium",
        "trigger": "旧会话清理失败后发起新连接",
        "system_result": "新实例继承旧会话状态",
        "external_observation": "新实例收到旧状态",
        "exclusion_condition": "旧会话已完全清理",
        "upstream_semantics": {
            "reachability": "重连入口可达",
            "caller_constraints": "调用方未阻断该路径",
            "documented_behavior": "REQ-HIST-01 要求新实例不接收旧状态",
            "existing_tests": "当前无确认回归用例",
            "conclusion": "risk_remains",
        },
        "translation_status": "Blackbox-ready",
        "status": "pending",
        "evidence": [historical_evidence],
    }]
    risk_result["analysis_checkpoint"]["failure_paths"][0].update({
        "linked_risk_ids": ["U00-R-HISTORY"],
        "failure": "旧会话清理不完整",
        "final_states": "新实例保留旧会话状态",
        "disposition": "risk",
    })
    write_json(Path(risk_task["result_path"]), risk_result)
    rejected = subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "validate-worker-result",
            "--task",
            str(risk_task_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "必须同时引用当前源码证据" in rejected.stderr

    risk_result["risks"][0]["evidence"].append(risk_result["evidence"][0])
    write_json(Path(risk_task["result_path"]), risk_result)
    _validate_worker_task(risk_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_TEST_GENERATION"

    test_task_path = Path(state["agent_task_paths"][0])
    test_task = read_json(test_task_path)
    test_result = _task_result(test_task)
    test_result["evidence"] = risk_result["evidence"]
    test_result["business_flows"] = risk_result["business_flows"]
    test_result["risks"] = risk_result["risks"]
    test_result["analysis_checkpoint"] = risk_result["analysis_checkpoint"]
    test_result["analysis_checkpoint"]["counterexamples_checked"] = [
        "历史故障只作失败观测，通过标准来自 REQ-HIST-01。",
    ]
    case = _test_case(
        "U00-TC-HISTORY",
        risk_ids=["U00-R-HISTORY"],
        requirement_ids=["REQ-HIST-01"],
        material_ids=[f"MAT:{historical_path}", "MAT:inbox/current.md"],
        title="重连清理历史问题回归",
    )
    case["steps"][0] = {
        "action": "完成旧会话清理后发起新连接",
        "expected_result": "新连接实例不接收旧会话状态",
        "failure_observation": "新连接实例收到旧会话状态",
    }
    test_result["test_cases"] = [case]
    write_json(Path(test_task["result_path"]), test_result)
    _validate_worker_task(test_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"
    markdown = Path(state["report_path"]).read_text(encoding="utf-8")
    html_report = Path(state["html_report_path"]).read_text(encoding="utf-8")
    for expected in (
        historical_path,
        "重连清理残留",
        f"MAT:{historical_path}",
        "MAT:inbox/current.md",
    ):
        assert expected in markdown and expected in html_report


def _multi_repo_isolation() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    state = run_module_analysis(str(contract))
    assert [unit["repo_id"] for unit in state["analysis_units"]] == ["repo-a", "repo-b"]
    tasks = [read_json(Path(path)) for path in state["agent_task_paths"]]
    assert len(tasks) == 2
    assert all(len(task["repositories"]) == 1 for task in tasks)
    assert {task["repositories"][0]["repo_id"] for task in tasks} == {"repo-a", "repo-b"}
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    review_task = read_json(
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review-independent.json"
    )
    assert {repo["repo_id"] for repo in review_task["repositories"]} == {"repo-a", "repo-b"}


def _unchanged_result_edit_does_not_block() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    state = run_module_analysis(str(contract))
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="REWORK")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK"
    run_dir = data_root / "runs" / "smoke-01"
    unchanged_path = run_dir / "agent-results" / "analysis" / "U01.json"
    original = read_json(unchanged_path)
    changed = dict(original)
    changed["summary"] = "初审后被修改"
    write_json(unchanged_path, changed)
    rework_task_path = run_dir / "agent-tasks" / "rework" / "U00.json"
    rework_task = read_json(rework_task_path)
    write_json(Path(rework_task["result_path"]), _task_result(rework_task))
    _validate_worker_task(rework_task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_REWORK_VERIFICATION"


def _run_ids_are_reserved_atomically() -> None:
    root = Path(tempfile.mkdtemp(prefix="pangea-run-id-smoke-"))
    _TEMP_ROOTS.append(root)
    contract = {"target": "CHAP", "data_root": str(root / "data")}
    with ThreadPoolExecutor(max_workers=8) as executor:
        run_ids = list(executor.map(lambda _: _reserve_run_id(contract), range(16)))
    assert len(set(run_ids)) == 16
    assert all((root / "data" / "runs" / run_id).is_dir() for run_id in run_ids)


def _module_analysis_owns_pending_contract_cleanup() -> None:
    _, _, contract = _workspace()
    pending = Path("pangea-data/.pangea/pending-task-contract-smoke-owned.json")
    pending.parent.mkdir(parents=True, exist_ok=True)
    write_json(pending, read_json(contract))
    start_module_analysis(str(pending))
    assert not pending.exists()

    start_module_analysis(str(contract))
    assert contract.exists()


def _agent_paths_and_contract_scope_are_posix() -> None:
    assert agent_path(Path(r"agent-tasks\analysis\U00.json")) == "agent-tasks/analysis/U00.json"
    base = {
        "run_id": "scope-path-smoke",
        "repository": "demo",
        "target": "CHAP",
    }
    for invalid_scope in (r"module\entry.c", "module\nentry.c"):
        try:
            TaskContract.model_validate({**base, "source_scope": [invalid_scope]})
        except ValueError:
            pass
        else:
            raise AssertionError(f"非法 source_scope 未被拒绝：{invalid_scope!r}")
    valid = TaskContract.model_validate({**base, "source_scope": ["module/entry.c"]})
    assert valid.source_scope == ["module/entry.c"]


def _directory_scope_splits_file_families_and_adds_unique_direct_callee() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    module = repo / "module"
    (module / "helper.h").write_text("int tls_handshake(void);\n", encoding="utf-8")
    (module / "helper.c").write_text(
        '#include "helper.h"\nint tls_handshake(void) { return 7; }\n',
        encoding="utf-8",
    )
    (module / "entry.c").write_text(
        '#include "helper.h"\nint demo_start(void) { return tls_handshake(); }\n',
        encoding="utf-8",
    )
    state = run_module_analysis(str(contract))
    tasks = [read_json(Path(path)) for path in state["agent_task_paths"]]
    assert len(tasks) == 1
    source_scopes = [set(task["unit"]["source_scope"]) for task in tasks]
    assert {frozenset(scope) for scope in source_scopes} == {
        frozenset({"module/entry.c", "module/helper.c", "module/helper.h"}),
    }
    assert len(set().union(*source_scopes)) == sum(len(scope) for scope in source_scopes)
    entry_task = next(task for task in tasks if "module/entry.c" in task["unit"]["source_scope"])
    assert "module/helper.c" not in entry_task["unit"]["context_scope"]
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="PASS")
    complete = run_module_analysis(str(contract))
    summaries = {item["unit_id"]: item for item in complete["analysis_summaries"]}
    entry_summary = summaries[entry_task["unit"]["unit_id"]]
    assert entry_summary["assigned_source_files"] == 3
    assert entry_summary["reviewed_source_files"] == 3
    assert entry_summary["function_count"] == 2
    assert entry_summary["direct_callee_context_count"] == 0
    assert "分析单元规模" in Path(complete["report_path"]).read_text(encoding="utf-8")


def _large_distinct_families_do_not_merge_into_one_unit() -> None:
    groups = [
        {
            "repo_id": "demo",
            "requested_scope": ["module/auth.c"],
            "code_paths": ["module/auth.c"],
            "context_paths": ["module/transport.c"],
        },
        {
            "repo_id": "demo",
            "requested_scope": ["module/transport.c"],
            "code_paths": ["module/transport.c"],
            "context_paths": [],
        },
    ]
    inventory = {
        "files": [
            {
                "repo_id": "demo",
                "path": "module/auth.c",
                "line_count": 1300,
                "functions": [{"symbol": f"auth_{index}"} for index in range(30)],
            },
            {
                "repo_id": "demo",
                "path": "module/transport.c",
                "line_count": 3153,
                "functions": [{"symbol": f"transport_{index}"} for index in range(90)],
            },
        ]
    }
    clustered = _cluster_groups(
        groups,
        inventory,
        [{
            "repo_id": "demo",
            "path": "module/transport.c",
            "reason": "direct_callee:transport_start",
        }],
    )
    assert [item["code_paths"] for item in clustered] == [
        ["module/auth.c"],
        ["module/transport.c"],
    ]


def _direct_caller_context_is_bounded() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.c"]
    write_json(contract, payload)
    callers = repo / "callers"
    callers.mkdir()
    for index in range(MAX_DIRECT_CALLERS_PER_GROUP + 10):
        (callers / f"caller_{index:02d}.c").write_text(
            "int demo_start(void);\n"
            f"int caller_{index:02d}(void) {{ return demo_start(); }}\n",
            encoding="utf-8",
        )
    (callers / "oversized.c").write_text(
        "\n".join([*("/* context */" for _ in range(900)), "int demo_start(void);", "int oversized(void) { return demo_start(); }"]),
        encoding="utf-8",
    )
    state = run_module_analysis(str(contract))
    direct_callers = [
        item
        for item in state["scope_expansion"]["context_files"]
        if str(item.get("reason", "")).startswith("direct_caller:")
    ]
    assert len(direct_callers) == MAX_DIRECT_CALLERS_PER_GROUP
    assert not any(item["path"].endswith("oversized.c") for item in direct_callers)


def _preprocessor_macro_is_not_a_function_definition() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "stubs.c").write_text(
        "#ifndef HAVE_REAL_HOOK\n"
        "LOG_REGISTER_COMPONENT(demo)\n"
        "#endif\n"
        "#ifndef HAVE_REAL_HOOK\n"
        "void\n"
        "demo_init_hooks(void)\n"
        "{\n"
        "}\n"
        "#endif\n",
        encoding="utf-8",
    )
    (repo / "module" / "real.c").write_text(
        "void demo_init_hooks(void)\n"
        "{\n"
        "}\n",
        encoding="utf-8",
    )
    (repo / "module" / "macro_user.c").write_text(
        "void macro_user(void) { LOG_REGISTER_COMPONENT(other); }\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/stubs.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "module/real.c" not in task["unit"]["context_scope"]
    assert "module/macro_user.c" not in task["unit"]["context_scope"]


def _multiline_condition_is_not_a_function_definition() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "common.c").write_text(
        "int demo_unlikely(int value);\n"
        "int common_entry(void) { return demo_unlikely(1); }\n",
        encoding="utf-8",
    )
    (repo / "app").mkdir()
    (repo / "app" / "reactor.c").write_text(
        "void unrelated(void)\n"
        "{\n"
        "    if (\n"
        "        demo_unlikely(1)) {\n"
        "    }\n"
        "}\n",
        encoding="utf-8",
    )
    (repo / "scripts").mkdir()
    (repo / "scripts" / "common.sh").write_text(
        "common common common common common common common common common common\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/common.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "app/reactor.c" not in task["unit"]["context_scope"]
    assert "scripts/common.sh" not in task["unit"]["context_scope"]


def _block_macro_is_not_an_exported_symbol() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "entry.c").write_text(
        "int demo_start(void)\n"
        "{\n"
        "    TAILQ_FOREACH(item, &items, link) {\n"
        "    }\n"
        "    return 0;\n"
        "}\n",
        encoding="utf-8",
    )
    (repo / "docs").mkdir()
    (repo / "docs" / "queues.md").write_text(
        "TAILQ_FOREACH(item, queue, link) is a generic queue macro.\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "docs/queues.md" not in task["unit"]["context_scope"]


def _oversized_direct_callee_is_not_frozen() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "entry.c").write_text(
        "int external_helper(void);\n"
        "int demo_start(void) { return external_helper(); }\n",
        encoding="utf-8",
    )
    (repo / "module" / "large_impl.c").write_text(
        "\n".join([
            *("/* implementation context */" for _ in range(900)),
            "int external_helper(void) { return 0; }",
        ]),
        encoding="utf-8",
    )
    (repo / "include").mkdir()
    (repo / "include" / "external.h").write_text(
        "int external_helper(void);\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "module/large_impl.c" not in task["unit"]["context_scope"]
    assert "include/external.h" in task["unit"]["context_scope"]
    assert any(
        item.get("path") == "include/external.h"
        and item.get("reason") == "oversized_callee_contract:external_helper"
        for item in state["scope_expansion"]["context_files"]
    )
    assert any(
        item.get("path") == "module/large_impl.c"
        and item.get("reason") == "direct_callee_context_too_large"
        for item in state["scope_expansion"]["unresolved_dependencies"]
    )


def _specialized_rubrics_follow_source_scope() -> None:
    cases = {
        "lib/iscsi/login.c": "storage_iscsi.md",
        "lib/nvmf/transport.c": "storage_nvmeof.md",
        "lib/env_dpdk/memory.c": "vendor_dpdk.md",
        "lib/rdma_utils/rdma_utils.c": "vendor_mlx_rdma.md",
        "lib/doca/task.c": "vendor_nvidia_doca.md",
    }
    for source_path, expected_name in cases.items():
        paths = checkpoint_rubrics(
            ["c_cpp"],
            [],
            repo_id="demo",
            source_paths=[source_path],
            inventory={"files": []},
        )
        assert any(path.endswith(expected_name) for path in paths), paths

    recovery_paths = checkpoint_rubrics(
        ["c_cpp"],
        [],
        repo_id="demo",
        source_paths=["module/session.c"],
        inventory={
            "files": [{
                "repo_id": "demo",
                "path": "module/session.c",
                "resource_signals": [{"keywords": ["alloc", "free"]}],
            }]
        },
    )
    assert any(path.endswith("storage_resource_recovery.md") for path in recovery_paths)

    unrelated = checkpoint_rubrics(
        ["c_cpp"],
        [],
        repo_id="demo",
        source_paths=["module/feature.c"],
        inventory={"files": []},
    )
    assert unrelated == ["src/pangea_agent/rubrics/builtin/c_cpp_analysis.md"]


def _doca_submit_cleanup_becomes_a_semantic_check() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    source = repo / "module" / "doca_task.c"
    source.write_text(
        "int doca_task_alloc_init(void *ctx, void **task);\n"
        "int doca_task_submit(void *task);\n"
        "int run(void *ctx) {\n"
        "  void *task = 0;\n"
        "  if (doca_task_alloc_init(ctx, &task) != 0) return -1;\n"
        "  if (doca_task_submit(task) != 0) return -2;\n"
        "  return 0;\n"
        "}\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/doca_task.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    checks = task["semantic_check_items"]
    assert len(checks) == 1
    assert checks[0]["check_id"] == "DOCA-SUBMIT-01"
    assert checks[0]["subject_path"] == "module/doca_task.c"
    assert "task free" in checks[0]["instruction"]


def _explicit_scope_keeps_external_implementation_as_context() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "api.h").write_text("int demo_external_api(void);\n", encoding="utf-8")
    (repo / "test").mkdir()
    (repo / "test" / "api.c").write_text(
        '#include "../module/api.h"\nint demo_external_api(void) { return 1; }\n',
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/api.h"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert task["unit"]["source_scope"] == ["module/api.h"]
    assert "test/api.c" in task["unit"]["context_scope"]


def _bounded_scope_expansion() -> None:
    _, _, contract = _workspace()
    repo = Path(read_json(contract)["data_root"]) / "repositories" / "demo"
    (repo / "module" / "demo_internal.h").write_text(
        "static inline int demo_abort(void) { assert(false); return 0; }\n"
        "static inline void demo_remove(void) { assert(STAILQ_EMPTY(&recv_stream)); }\n"
        "static inline void demo_release(void) { assert(entry->ref > 0); }\n"
        "static inline void demo_state(void) { assert(sock->pipe_has_data == false); }\n"
        "static inline void demo_mark(void) { sock->pipe_has_data = true; }\n"
        "static inline void demo_set_pipe(void) { destroy_pipe(); }\n",
        encoding="utf-8",
    )
    (repo / "module" / "unused_internal.h").write_text(
        "static inline int demo_unused(void) { return 0; }\n",
        encoding="utf-8",
    )
    (repo / "module" / "entry.c").write_text(
        '#include "demo_internal.h"\n#include "unused_internal.h"\n'
        "int demo_feature_start(void) { return demo_abort(); }\n"
        "int demo_add(void) { return spdk_sock_map_insert(); }\n"
        "void demo_drop(void) { spdk_sock_map_release(); }\n",
        encoding="utf-8",
    )
    (repo / "app").mkdir()
    (repo / "app" / "rpc.c").write_text(
        "int demo_feature_start(void);\n"
        "int rpc_start(void) { return demo_feature_start(); }\n"
        "int rpc_add(void) { return spdk_sock_map_insert(); }\n"
        "void rpc_drop(void) { spdk_sock_map_release(); }\n",
        encoding="utf-8",
    )
    (repo / "unrelated").mkdir()
    (repo / "unrelated" / "noise.c").write_text("int unrelated(void) { return 0; }\n", encoding="utf-8")
    (repo / "test" / "e2e").mkdir(parents=True)
    (repo / "test" / "e2e" / "demo.sh").write_text("demo_feature_start()\n", encoding="utf-8")
    (repo / "docs").mkdir()
    (repo / "docs" / "demo.md").write_text(
        "demo feature context " * 20,
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["target"] = "demo feature"
    payload["source_scope"] = ["module/entry.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "module/entry.c" in task["unit"]["source_scope"]
    assert "module/demo_internal.h" in task["unit"]["context_scope"]
    assert "module/unused_internal.h" not in task["unit"]["context_scope"]
    assert "docs/demo.md" not in task["unit"]["context_scope"]
    assert task["failure_signal_context"] == [
        {
            "path": "module/demo_internal.h",
            "line": 1,
            "signal": "static inline int demo_abort(void) { assert(false); return 0; }",
            "analysis_focus": (
                "先定位直接支配 assert 的失败条件，再分别重放 Debug 与 Release。受支持模式中的底层"
                "操作若可返回失败，且公开契约或入口没有阻断，Debug 终止必须保留为风险；不能用"
                " assert 后的清理或返回排除，Release 继续核对清理后的最终状态。条件含数值句柄时，"
                "必须从创建函数的失败返回值确认哨兵，并把 0 作为独立边界重放。"
            ),
            "related_state_context": [],
        },
        {
            "path": "module/demo_internal.h",
            "line": 2,
            "signal": "static inline void demo_remove(void) { assert(STAILQ_EMPTY(&recv_stream)); }",
            "analysis_focus": (
                "追踪容器元素的产生、归还和公开移除入口；实现注释或 assert 本身不是调用方契约，"
                "只有公开契约或入口强制检查才能证明该状态不可达。"
            ),
            "related_state_context": [],
        },
        {
            "path": "module/demo_internal.h",
            "line": 3,
            "signal": "static inline void demo_release(void) { assert(entry->ref > 0); }",
            "analysis_focus": (
                "按任务提供的每个直接实现写出实际的增加与减少调用序列，并追踪错误日志之后的函数"
                "返回值和上层是否真正绑定对象。只有证明某次 release/decrement 前没有成功 insert/"
                "increment 才能判定失衡；lookup 不增加引用本身不足以证明风险。一个实现的结论不能"
                "覆盖另一个实现。"
            ),
            "related_state_context": [],
        },
        {
            "path": "module/demo_internal.h",
            "line": 4,
            "signal": "static inline void demo_state(void) { assert(sock->pipe_has_data == false); }",
            "analysis_focus": (
                "把断言可达性与重配置后的状态残留拆成两条 failure path。先判断断言本身，再从状态置位"
                "重放 related_state_context 中的 destroy/NULL/setter；即使断言不可达，重配置仍可能独立"
                "造成数据丢失或残留状态。当前分支没有写入者不能证明先前状态不会残留。"
            ),
            "related_state_context": [
                "module/demo_internal.h:5: static inline void demo_mark(void) { sock->pipe_has_data = true; }",
                "module/demo_internal.h:6: static inline void demo_set_pipe(void) { destroy_pipe(); }",
            ],
        },
    ]
    assert "app/rpc.c" in task["unit"]["context_scope"]
    assert "unrelated/noise.c" not in task["unit"]["source_scope"]
    assert "test/e2e/demo.sh" in task["unit"]["context_scope"]
    assert len(task["unit"]["context_scope"]) <= 10
    assert task["max_parallel_workers"] == 8 and task["may_spawn_workers"] is False
    semantic_checks = task["semantic_check_items"]
    semantic_kinds = [item["kind"] for item in semantic_checks]
    assert semantic_kinds == [
        "assertion_reachability",
        "assertion_reachability",
        "paired_operation",
        "assertion_reachability",
        "resource_reconfiguration",
        "resource_reconfiguration",
    ], semantic_kinds
    assert [item["subject_path"] for item in semantic_checks[2:3]] == [
        "module/entry.c",
    ]
    assert len({item["check_id"] for item in semantic_checks}) == len(semantic_checks)


def _state_context_balances_lifecycle_and_reconfiguration() -> None:
    lines = [
        *(f"sock->pending_recv = {'true' if index % 2 else 'false'};" for index in range(8)),
        *(f"set_recv_pipe_{index}();" for index in range(8)),
    ]
    context = _related_state_context(
        "module/demo.c", lines, "assert(sock->pending_recv == false);"
    )
    assert len(context) == 12
    assert context[:3] == [
        "module/demo.c:1: sock->pending_recv = false;",
        "module/demo.c:2: sock->pending_recv = true;",
        "module/demo.c:3: sock->pending_recv = false;",
    ]
    assert context[3:6] == [
        "module/demo.c:6: sock->pending_recv = true;",
        "module/demo.c:7: sock->pending_recv = false;",
        "module/demo.c:8: sock->pending_recv = true;",
    ]
    assert context[6] == "module/demo.c:9: set_recv_pipe_0();"
    assert context[-1] == "module/demo.c:16: set_recv_pipe_7();"


def _selected_header_keeps_its_direct_local_dependency() -> None:
    _, data_root, contract = _workspace()
    repo = data_root / "repositories" / "demo"
    (repo / "module" / "entry.c").write_text(
        '#include "demo_internal.h"\nint demo_start(void) { return demo_config(); }\n',
        encoding="utf-8",
    )
    (repo / "module" / "demo_internal.h").write_text(
        '#include "demo_spec.h"\n'
        "static inline int demo_config(void) { return DEMO_QUEUE_LIMIT; }\n",
        encoding="utf-8",
    )
    (repo / "module" / "demo_spec.h").write_text(
        "enum demo_limits { DEMO_QUEUE_LIMIT = 16 };\n",
        encoding="utf-8",
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.c"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert "module/demo_internal.h" in task["unit"]["context_scope"]
    assert "module/demo_spec.h" in task["unit"]["context_scope"]
    assert any(
        item.get("path") == "module/demo_spec.h"
        and item.get("reason") == "direct_header_dependency:module/demo_internal.h"
        for item in state["scope_expansion"]["context_files"]
    )


def _expected_behavior_not_risk() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    result = _task_result(task)
    result["risks"] = [{
        "risk_id": "R-EXPECTED",
        "title": "规格已经定义的行为",
        "affected_paths": task["unit"]["source_scope"],
        "dfx": ["功能与状态"],
        "severity": "Low",
        "confidence": "high",
        "trigger": "触发条件",
        "system_result": "规格行为",
        "external_observation": "可观测",
        "exclusion_condition": "无",
        "upstream_semantics": {
            "reachability": "可达",
            "caller_constraints": "调用方按规格处理",
            "documented_behavior": "高层 API 明确定义为预期行为",
            "existing_tests": "已有测试验证该规格行为",
            "conclusion": "expected_behavior",
        },
        "translation_status": "Blackbox-ready",
        "status": "pending",
        "evidence": result["evidence"],
    }]
    result["test_cases"] = [_test_case("TC-EXPECTED", risk_ids=["R-EXPECTED"])]
    result["analysis_checkpoint"]["failure_paths"][0]["disposition"] = "risk"
    result["analysis_checkpoint"]["failure_paths"][0]["linked_risk_ids"] = ["R-EXPECTED"]
    write_json(Path(task["result_path"]), result)
    assert "预期行为" in _reject_worker_task(Path(state["agent_task_paths"][0]))


def _unversioned_source_is_deliverable() -> None:
    _, data_root, contract = _workspace()
    shutil.rmtree(data_root / "repositories" / "demo" / ".git")
    state = run_module_analysis(str(contract))
    state = _advance_to_test_generation(state)
    _write_all_analysis(state)
    run_module_analysis(str(contract))
    _review(data_root, "smoke-01", status="PASS")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"
    assert not any(
        item.get("kind") == "source_version_unverifiable"
        for item in state["quality_report"]["unresolved"]
    )


def _known_c_macro_parse_artifacts() -> None:
    assert _known_macro_parse_artifact(
        {"line": 1, "text": ","}, ["TAILQ_HEAD(, item) head;"]
    )
    assert _known_macro_parse_artifact(
        {"line": 2, "text": "cc."}, ["return offsetof(struct registers,", "cc.raw);"]
    )
    assert _known_macro_parse_artifact(
        {"line": 2, "text": "struct"}, ["ctx =", "SPDK_CONTAINEROF(arg, struct ctx, member);"]
    )
    assert _known_macro_parse_artifact(
        {"line": 1, "text": "hex_char"}, ["static const char __spdk_nonstring hex_char[16];"]
    )
    assert _known_macro_parse_artifact(
        {"line": 3, "text": "void"},
        ["SPDK_LOG_REGISTER_COMPONENT(demo)", "", "static void"],
    )
    assert _known_macro_parse_artifact(
        {"line": 2, "text": "void"},
        ["callback(Client *client,", "    AVAHI_GCC_UNUSED void *user_data)"],
    )
    assert not _known_macro_parse_artifact(
        {"line": 1, "text": "unexpected"}, ["int broken = ;"]
    )


def _rework_review_prunes_removed_links() -> None:
    root = Path(tempfile.mkdtemp(prefix="pangea-review-links-"))
    _TEMP_ROOTS.append(root)
    result_path = root / "worker.json"
    manifest_path = root / "manifest.json"
    write_json(manifest_path, {"material_catalog": []})
    worker_task = {
        "run_id": "smoke-links",
        "unit": {
            "unit_id": "U00",
            "repo_id": "demo",
            "source_scope": ["module/entry.c"],
            "context_scope": [],
        },
        "stage": "rework",
        "attempt": 1,
        "review_issues": [],
        "semantic_check_items": [],
        "source_manifest_path": str(manifest_path),
    }
    write_json(result_path, _task_result(worker_task))
    independent = IndependentReviewResult.model_validate({
        "run_id": "smoke-links",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "summary": "独立复核完成",
        "reviewed_units": ["U00"],
        "findings": [{
            "unit_id": "U00",
            "check_id": "CHECK-1",
            "finding": "独立检查结果",
            "evidence": ["demo:module/entry.c:1"],
        }],
    })
    prior = ReviewResult.model_validate({
        "run_id": "smoke-links",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "REWORK",
        "summary": "需要返工",
        "issues": [{
            "issue_id": "I-1",
            "unit_id": "U00",
            "reason": "删除旧风险",
            "required_change": "删除旧风险和用例",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [{
            **independent.findings[0].model_dump(mode="json"),
            "worker_disposition": "contradiction",
            "linked_worker_risk_ids": ["REMOVED-RISK"],
            "linked_worker_test_case_ids": ["REMOVED-TEST"],
        }],
        "test_case_checks": [],
    })
    task = ReviewTask.model_validate({
        "run_id": "smoke-links",
        "target": "link pruning",
        "repositories": [{"repo_id": "demo", "source_root": str(root)}],
        "inventory_path": str(root / "inventory.json"),
        "source_manifest_path": str(root / "manifest.json"),
        "stage": "rework_verification",
        "result_path": str(root / "review.json"),
        "analysis_tasks": [{"unit_id": "U00", "task_path": str(root / "task.json")}],
        "analysis_results": [{"unit_id": "U00", "result_path": str(result_path)}],
        "independent_result_path": str(root / "independent.json"),
        "same_reviewer_id": "reviewer-1",
        "prior_issues": [prior.issues[0].model_dump(mode="json")],
        "may_spawn_workers": False,
        "review_round": 1,
    })
    skeleton = review_result_skeleton(task, independent, prior)
    finding = skeleton["independent_findings"][0]
    assert finding["linked_worker_risk_ids"] == []
    assert finding["linked_worker_test_case_ids"] == []


def _review_issue_requires_actionable_change() -> None:
    assert _finding_excludes_linked_leak(
        "返回失败后调用 destruct，所有 qbase_addr 全部释放，未发现资源泄漏",
        "DMA channel initialization memory leak",
    )
    assert not _finding_excludes_linked_leak(
        "返回失败后仅释放 ring，qbase_addr 仍保留",
        "DMA channel initialization memory leak",
    )
    issues = ReviewResult.model_validate({
        "run_id": "smoke-actionable",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "UNRESOLVED",
        "summary": "仍有问题",
        "issues": [{
            "issue_id": "I-VAGUE",
            "unit_id": "U00",
            "reason": "TC-1 step-level logical contradiction persists after rework.",
            "required_change": "Consider updating TC-1.",
        }, {
            "issue_id": "I-DEFINITE",
            "unit_id": "U00",
            "reason": "旧步骤与已确认源码分支相反",
            "required_change": "考虑到源码分支只返回 -1，删除 TC-2 中返回 0 的步骤",
        }, {
            "issue_id": "I-ACTIONABLE-SUGGESTION",
            "unit_id": "U00",
            "reason": "TC-3 仍引用旧返回值",
            "required_change": "建议修改 TC-3 第 2 步 expected_result 为返回 -1",
        }, {
            "issue_id": "I-NO-REWORK",
            "unit_id": "U00",
            "reason": "该 finding 属于范围外依赖",
            "required_change": "确认现有 Developer-confirm 结论正确，无需返工。",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": [],
    }).issues
    assert _non_actionable_review_issue_ids(issues) == ["I-VAGUE", "I-NO-REWORK"]
    reviewer_owned = ReviewResult.model_validate({
        "run_id": "smoke-reviewer-owned",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "UNRESOLVED",
        "summary": "Reviewer 字段未填",
        "issues": [{
            "issue_id": "I-REVIEWER-OWNED",
            "unit_id": "U00",
            "reason": "current_behavior 仍为占位符",
            "required_change": "填写 test_case_checks 的 current_behavior 字段",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": [],
    }).issues
    assert _reviewer_owned_field_issue_ids(reviewer_owned) == ["I-REVIEWER-OWNED"]

    reviewer_correction = ReviewResult.model_validate({
        "run_id": "smoke-reviewer-correction",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "UNRESOLVED",
        "summary": "Reviewer 自身结论需纠正",
        "issues": [{
            "issue_id": "I-INDEPENDENT-CORRECTION",
            "unit_id": "U00",
            "reason": "independent finding 与冻结源码不符",
            "required_change": "将 independent finding 的 worker_disposition 改为 contradiction",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": [],
    }).issues
    assert _reviewer_self_correction_issue_ids(reviewer_correction) == [
        "I-INDEPENDENT-CORRECTION"
    ]

    reviewer_scope_excuse = ReviewResult.model_validate({
        "run_id": "smoke-reviewer-scope-excuse",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "UNRESOLVED",
        "summary": "Reviewer 尝试把自身范围错误写成 issue",
        "issues": [{
            "issue_id": "I-REVIEWER-SCOPE",
            "unit_id": "U00",
            "reason": "Reviewer self-correction: finding 的 scope 在原始 review 有问题",
            "required_change": "无需 worker 返工，后续 reviewer 自行修正",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": [],
    }).issues
    assert _reviewer_self_correction_issue_ids(reviewer_scope_excuse) == [
        "I-REVIEWER-SCOPE"
    ]

    stale_restore = ReviewResult.model_validate({
        "run_id": "smoke-stale-restore",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "UNRESOLVED",
        "summary": "错误恢复已删除用例",
        "issues": [{
            "issue_id": "I-STALE-RESTORE",
            "unit_id": "U00",
            "reason": "旧 TC04 已从当前结果删除",
            "required_change": "恢复 TC04 并标记 invalid",
        }],
        "reviewed_units": ["U00"],
        "independent_findings": [],
        "test_case_checks": [],
    }).issues
    assert _stale_artifact_restoration_issue_ids(
        stale_restore,
        {("U00", "R00")},
        {("U00", "TC00")},
    ) == ["I-STALE-RESTORE"]


def _literal_boundary_comparison_cannot_be_reversed() -> None:
    _, _, contract = _workspace()
    state = _advance_to_test_generation(run_module_analysis(str(contract)))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    result = _task_result(task)
    result["summary"] = "offset=0, size=0，因此 0 >= 0 evaluates false"
    write_json(Path(task["result_path"]), result)
    assert "源码边界判断存在可直接求值的事实错误" in _reject_worker_task(task_path)
    result["summary"] = "offset=0, size=0，因此 0 >= 0 evaluates true"
    result["test_cases"] = [_test_case(
        "TC-ASAN-ORACLE", requirement_ids=["REQ-ASAN-ORACLE"]
    )]
    result["test_cases"][0]["steps"][0]["expected_result"] = "ASan 报告 double-free"
    write_json(Path(task["result_path"]), result)
    assert "缺陷信号只能写 failure_observation" in _reject_worker_task(task_path)
    result["test_cases"][0]["steps"][0]["expected_result"] = (
        "ASan 不应报告 double-free，进程不崩溃且正常退出"
    )
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(task_path)


def _source_checkpoint_uses_frozen_inputs() -> None:
    _, data_root, contract = _workspace()
    raw_contract = read_json(contract)
    state = load_contract({"run_id": "smoke-01", "data_root": str(data_root), "task_contract": raw_contract})
    state = locate_module(resolve_repositories(state))
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert progress["init_step"] == "SOURCE_READY"
    original = data_root / "repositories" / "demo" / "module" / "entry.c"
    original.write_text("int changed_after_checkpoint(void) { return 1; }\n", encoding="utf-8")
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    frozen = Path(task["repositories"][0]["source_root"]) / "module" / "entry.c"
    assert "demo_start" in frozen.read_text(encoding="utf-8")
    assert "changed_after_checkpoint" not in frozen.read_text(encoding="utf-8")


def _source_checkpoint_rejects_known_ae4dma_misreads() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    frozen = Path(task["repositories"][0]["source_root"]) / "module" / "entry.c"
    frozen.write_text(
        """int process_events(struct cmd_queue *cmd_q) {
    uint64_t sub_desc_cnt = cmd_q->ring_buff_count;
    while (sub_desc_cnt) {
        if (desc_status == AE4DMA_DMA_DESC_SUBMITTED) {
            break;
        }
        assert(cmd_q->ring_buff_count > 0);
        cmd_q->ring_buff_count--;
        if (cmd_q->ring[tail].callback_fn) {
            cmd_q->ring[tail].callback_fn(cmd_q->ring[tail].callback_arg, desc_err_code);
        }
        sub_desc_cnt--;
    }
}
void ae4dma_channel_destruct(struct chan *ae4dma) {
    ae4dma_unmap_pci_bar(ae4dma);
    spdk_free(ae4dma->cmd_q[0].qbase_addr);
    free(ae4dma->cmd_q[0].ring);
}
void *ae4dma_attach(struct chan *ae4dma) {
    if (ae4dma_channel_start(queues, ae4dma) != 0) {
        ae4dma_channel_destruct(queues, ae4dma);
        free(ae4dma);
        return NULL;
    }
}
""",
        encoding="utf-8",
    )
    result = _task_result(task)
    failure_path = result["analysis_checkpoint"]["failure_paths"][0]
    false_candidates = (
        "ring calloc 分配失败后 qbase_addr 未释放并发生资源泄漏",
        "spdk_dma_zmalloc 分配失败后 PCI BAR 映射未释放并发生泄漏",
        "callback_fn 为 NULL 时仍被解引用，触发空指针崩溃",
        "遇到 AE4DMA_DMA_DESC_SUBMITTED 后 ring_buff_count 与 sub_desc_cnt 计数不一致，队列永久 stuck",
    )
    for claim in false_candidates:
        failure_path.update({
            "trigger": claim,
            "side_effects": "候选声称出现异常状态",
            "failure": claim,
            "caller_handling": "候选声称调用方无法恢复",
            "final_states": "候选声称资源或队列保持异常",
        })
        write_json(Path(task["result_path"]), result)
        rejection = _reject_worker_task(task_path)
        assert "不得把源码检查之前尚未发生" in rejection
        assert failure_path["path_id"] in rejection


def _index_checkpoint_resumes_without_live_source() -> None:
    _, data_root, contract = _workspace()
    raw_contract = read_json(contract)
    state = load_contract({"run_id": "smoke-01", "data_root": str(data_root), "task_contract": raw_contract})
    state = locate_module(resolve_repositories(state))
    index_materials(state)
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert progress["init_step"] == "INDEX_READY"
    (data_root / "repositories" / "demo" / "module" / "entry.c").unlink()
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert Path(state["agent_task_paths"][0]).is_file()


def _agent_start_checkpoint() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "prepare-worker-result", "--task", str(task_path)],
        check=True,
        capture_output=True,
    )
    result_path = Path(read_json(task_path)["result_path"])
    prepared_result = read_json(result_path)
    assert prepared_result["completed_stage"] == "source_checkpoint"
    assert prepared_result["analysis_checkpoint"]["source_paths_reviewed"] == []
    prepared_result["worker_id"] = "worker-empty-checkpoint"
    prepared_result["summary"] = "尚未填写源码 checkpoint"
    write_json(result_path, prepared_result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(task_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "lifecycle_stages_checked" in rejected.stderr
    source_task = read_json(task_path)
    assert source_task["inventory_path"] is None
    assert source_task["source_manifest_path"] is None
    assert source_task["coverage_context"] == []
    assert state["agent_actions"][0]["action"] == "dispatch_agent"
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "analysis", "--unit-id", "U00", "--task-id", _ANALYSIS_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    progress_path = data_root / "runs" / "smoke-01" / "progress.json"
    assert read_json(progress_path)["agent_sessions"]["analysis:U00"]["status"] == "dispatched"
    state = resume_module_analysis(
        "smoke-01",
        str(data_root),
        settled_task_id=_ANALYSIS_SESSION_ID,
    )
    assert state["agent_actions"][0]["action"] == "continue_agent"
    assert state["agent_actions"][0]["task_id"] == _ANALYSIS_SESSION_ID
    state = _advance_to_test_generation(state)
    active_task = read_json(Path(state["agent_task_paths"][0]))
    active_result = read_json(Path(active_task["result_path"]))
    assert active_result["worker_id"] == _ANALYSIS_SESSION_ID
    task = read_json(Path(state["agent_task_paths"][0]))
    write_json(Path(task["result_path"]), _task_result(task))
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    state = run_module_analysis(str(contract))
    review_path = data_root / "runs" / "smoke-01" / "agent-tasks" / "review-independent.json"
    subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "prepare-review-result", "--task", str(review_path)],
        check=True,
        capture_output=True,
    )
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "record-agent-session",
            "--run-id",
            "smoke-01",
            "--data-root",
            str(data_root),
            "--role",
            "review",
            "--task-id",
            _REVIEW_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    progress = read_json(progress_path)
    assert progress["agent_sessions"]["review"]["status"] == "dispatched"
    comparison_task, _ = _advance_to_review_comparison(data_root, "smoke-01")
    assert comparison_task["same_reviewer_id"] == _REVIEW_SESSION_ID
    progress = read_json(progress_path)
    assert progress["agent_sessions"]["review"] == {
        "role": "review",
        "unit_id": None,
        "stage": "comparison_review",
        "task_id": _REVIEW_SESSION_ID,
        "status": "pending",
    }
    comparison_task_path = (
        data_root / "runs" / "smoke-01" / "agent-tasks" / "review.json"
    )
    subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "prepare-review-result",
            "--task",
            str(comparison_task_path),
        ],
        check=True,
        capture_output=True,
    )
    prepared_review = read_json(
        data_root / "runs" / "smoke-01" / "agent-results" / "review.json"
    )
    assert prepared_review["reviewer_id"] == _REVIEW_SESSION_ID
    assert read_json(progress_path)["agent_sessions"]["review"]["status"] == "pending"

    prepared_review["summary"] = "尚未通过的损坏内容"
    prepared_review["independent_findings"][0]["finding"] = "被错误改写"
    write_json(
        data_root / "runs" / "smoke-01" / "agent-results" / "review.json",
        prepared_review,
    )
    subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "prepare-review-result",
            "--task",
            str(comparison_task_path),
            "--fresh",
        ],
        check=True,
        capture_output=True,
    )
    rebuilt_review = read_json(
        data_root / "runs" / "smoke-01" / "agent-results" / "review.json"
    )
    independent_result = read_json(
        data_root / "runs" / "smoke-01" / "agent-results" / "review-independent.json"
    )
    assert rebuilt_review["summary"] == ""
    assert rebuilt_review["issues"] == []
    assert rebuilt_review["independent_findings"][0]["finding"] == (
        independent_result["findings"][0]["finding"]
    )


def _record_agent_session_rejects_fake_task_id() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    progress_path = data_root / "runs" / "smoke-01" / "progress.json"
    before = read_json(progress_path)

    for fake_task_id in ("null", "SUBAGENT_PLACEHOLDER"):
        rejected = subprocess.run(
            [
                sys.executable,
                "-m",
                "pangea_agent.cli.main",
                "record-agent-session",
                "--run-id",
                "smoke-01",
                "--data-root",
                str(data_root),
                "--role",
                "analysis",
                "--unit-id",
                "U00",
                "--task-id",
                fake_task_id,
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert rejected.returncode != 0
        assert "真实 subagent UUID" in rejected.stderr
        assert read_json(progress_path) == before
    root_completion = subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "record-agent-session",
            "--run-id",
            "smoke-01",
            "--data-root",
            str(data_root),
            "--role",
            "analysis",
            "--unit-id",
            "U00",
            "--task-id",
            _ANALYSIS_SESSION_ID,
            "--status",
            "completed",
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert root_completion.returncode != 0
    assert "unrecognized arguments" in root_completion.stderr
    assert read_json(progress_path) == before
    assert state["agent_actions"][0]["action"] == "dispatch_agent"
    assert state["agent_actions"][0]["task_id"] is None


def _parallel_worker_validation_preserves_completions() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    state = run_module_analysis(str(contract))
    task_paths = [Path(path) for path in state["agent_task_paths"]]
    session_ids = (_ANALYSIS_SESSION_ID, _REPLACEMENT_SESSION_ID)
    for action, task_path, session_id in zip(state["agent_actions"], task_paths, session_ids):
        subprocess.run(
            [
                sys.executable,
                "-m",
                "pangea_agent.cli.main",
                "record-agent-session",
                "--run-id",
                "smoke-01",
                "--data-root",
                str(data_root),
                "--role",
                "analysis",
                "--unit-id",
                action["unit_id"],
                "--task-id",
                session_id,
            ],
            check=True,
            capture_output=True,
        )
        task = read_json(task_path)
        write_json(Path(task["result_path"]), _task_result(task))

    processes = [
        subprocess.Popen(
            [
                sys.executable,
                "-m",
                "pangea_agent.cli.main",
                "validate-worker-result",
                "--task",
                str(task_path),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        for task_path in task_paths
    ]
    outputs = [process.communicate() for process in processes]
    assert all(process.returncode == 0 for process in processes), outputs
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert all(
        progress["agent_sessions"][f"analysis:U0{index}"]["status"] == "completed"
        for index in range(2)
    )
    assert run_module_analysis(str(contract))["phase"] == "WAITING_RISK_ANALYSIS"


def _fast_worker_completion_survives_late_session_binding() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    write_json(Path(task["result_path"]), _task_result(task))
    _validate_worker_task(task_path)

    progress_path = data_root / "runs" / "smoke-01" / "progress.json"
    session = read_json(progress_path)["agent_sessions"]["analysis:U00"]
    assert session["task_id"] is None and session["status"] == "completed"
    subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "record-agent-session",
            "--run-id",
            "smoke-01",
            "--data-root",
            str(data_root),
            "--role",
            "analysis",
            "--unit-id",
            "U00",
            "--task-id",
            _ANALYSIS_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    session = read_json(progress_path)["agent_sessions"]["analysis:U00"]
    assert session["task_id"] == _ANALYSIS_SESSION_ID
    assert session["status"] == "completed"
    assert run_module_analysis(str(contract))["phase"] == "WAITING_RISK_ANALYSIS"


def _unvalidated_worker_returns_graph_continue_action() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "analysis", "--unit-id", "U00", "--task-id", _ANALYSIS_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "prepare-worker-result",
            "--task", str(task_path),
        ],
        check=True,
        capture_output=True,
    )

    resumed = resume_module_analysis(
        "smoke-01",
        str(data_root),
        settled_task_id=_ANALYSIS_SESSION_ID,
    )
    assert resumed["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert resumed["agent_actions"] == [{
        "action": "continue_agent",
        "role": "analysis",
        "stage": "source_checkpoint",
        "session_key": "analysis:U00",
        "unit_id": "U00",
        "task_path": str(task_path),
        "task_id": _ANALYSIS_SESSION_ID,
        "replacement_allowed": False,
        "after_completion": "resume_run",
    }]

    task = read_json(task_path)
    write_json(Path(task["result_path"]), _task_result(task))
    _validate_worker_task(task_path)
    assert run_module_analysis(str(contract))["phase"] == "WAITING_RISK_ANALYSIS"


def _nonexecuted_callback_cannot_write_state() -> None:
    _, _, contract = _lua_workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    result = _task_result(task)
    target = next(
        item for item in result["analysis_checkpoint"]["failure_paths"]
        if item["path_id"].startswith("SC-LUA-ERROR-")
    )
    target["side_effects"] = "C1 已注册但执行0次"
    target["final_states"] = "callback_count=1（C1 函数体写入）"
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(task_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "已声明未执行" in rejected.stderr

    target["final_states"] = "callback_count=0（C1 未执行，未写入）"
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(task_path)


def _graph_v2_stage_actions_reach_complete() -> None:
    _, data_root, contract = _workspace()
    run_dir = data_root / "runs" / "smoke-01"
    progress_path = run_dir / "progress.json"

    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert read_json(progress_path)["workflow_version"] == 2
    task_path = Path(state["agent_task_paths"][0])
    initial_task = read_json(task_path)
    assert initial_task["stage"] == "source_checkpoint"
    initial_result_path = Path(initial_task["result_path"])
    assert initial_result_path.is_file()
    initial_result = read_json(initial_result_path)
    assert initial_result["completed_stage"] == "source_checkpoint"
    assert initial_result["analysis_checkpoint"]["source_paths_reviewed"] == []
    assert initial_result["risks"] == []
    assert initial_result["test_cases"] == []
    assert state["agent_actions"] == [{
        "action": "dispatch_agent",
        "role": "analysis",
        "stage": "source_checkpoint",
        "session_key": "analysis:U00",
        "unit_id": "U00",
        "task_path": str(task_path),
        "task_id": None,
        "replacement_allowed": False,
        "after_completion": "resume_run",
    }]
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "analysis", "--unit-id", "U00", "--task-id", _ANALYSIS_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )

    task_paths_seen = {task_path}
    for completed_stage, next_phase, next_stage in (
        ("source_checkpoint", "WAITING_RISK_ANALYSIS", "risk_analysis"),
        ("risk_analysis", "WAITING_TEST_GENERATION", "test_generation"),
    ):
        task = read_json(task_path)
        result = _task_result(task)
        assert result["completed_stage"] == completed_stage
        write_json(Path(task["result_path"]), result)
        if completed_stage == "source_checkpoint":
            waiting = resume_module_analysis(
                "smoke-01",
                str(data_root),
                settled_task_id=_ANALYSIS_SESSION_ID,
            )
            assert waiting["phase"] == "WAITING_SOURCE_CHECKPOINT"
            assert waiting["agent_actions"][0]["action"] == "continue_agent"
            assert waiting["agent_actions"][0]["task_id"] == _ANALYSIS_SESSION_ID
            assert read_json(progress_path)["agent_sessions"]["analysis:U00"]["status"] == "pending"
        _validate_worker_task(task_path)
        state = run_module_analysis(str(contract))
        assert state["phase"] == next_phase
        assert read_json(task_path)["stage"] == completed_stage
        next_task_path = Path(state["agent_task_paths"][0])
        assert next_task_path not in task_paths_seen
        task_paths_seen.add(next_task_path)
        task_path = next_task_path
        assert read_json(task_path)["stage"] == next_stage
        if next_stage == "risk_analysis":
            risk_task = read_json(task_path)
            assert risk_task["inventory_path"]
            assert risk_task["source_manifest_path"]
        assert state["agent_actions"][0]["action"] == "continue_agent"
        assert state["agent_actions"][0]["task_id"] == _ANALYSIS_SESSION_ID

    task = read_json(task_path)
    result = _task_result(task)
    assert result["completed_stage"] == "test_generation"
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    independent_path = run_dir / "agent-tasks" / "review-independent.json"
    independent = read_json(independent_path)
    assert independent["stage"] == "independent_review"
    assert state["agent_actions"][0]["action"] == "dispatch_agent"
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
            "--run-id", "smoke-01", "--data-root", str(data_root),
            "--role", "review", "--task-id", _REVIEW_SESSION_ID,
        ],
        check=True,
        capture_output=True,
    )
    findings = []
    for reference in independent["analysis_tasks"]:
        worker_task = read_json(Path(reference["task_path"]))
        unit_findings = [{
            "unit_id": reference["unit_id"],
            "check_id": item["check_id"],
            "finding": f"独立核对 {item['check_id']} 未发现额外问题",
            "evidence": [f"{worker_task['unit']['repo_id']}:{item['subject_path']}:1"],
        } for item in worker_task["semantic_check_items"]]
        if not unit_findings:
            unit_findings.append({
                "unit_id": reference["unit_id"],
                "check_id": f"BASELINE-{reference['unit_id']}",
                "finding": "独立核对正常与异常路径，未发现额外问题",
                "evidence": [
                    f"{worker_task['unit']['repo_id']}:{worker_task['unit']['source_scope'][0]}:1"
                ],
            })
        findings.extend(unit_findings)
    write_json(Path(independent["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "summary": "独立复核完成",
        "reviewed_units": [item["unit_id"] for item in independent["analysis_tasks"]],
        "findings": findings,
    })
    _validate_review_task(independent_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_COMPARISON_REVIEW"
    comparison = read_json(run_dir / "agent-tasks" / "review.json")
    assert comparison["stage"] == "comparison_review"
    assert state["agent_actions"][0]["action"] == "continue_agent"
    assert state["agent_actions"][0]["task_id"] == _REVIEW_SESSION_ID
    write_json(Path(comparison["result_path"]), {
        "schema_version": "1.0",
        "run_id": "smoke-01",
        "reviewer_id": "reviewer-1",
        "finish_reason": "stop",
        "status": "PASS",
        "summary": "对照复核完成",
        "issues": [],
        "reviewed_units": [item["unit_id"] for item in comparison["analysis_results"]],
        "independent_findings": _compared_findings(comparison, findings),
        "test_case_checks": _test_case_checks(comparison),
    })
    _validate_review_task(run_dir / "agent-tasks" / "review.json")
    state = run_module_analysis(str(contract))
    assert state["phase"] == "COMPLETE"


def _graph_v2_two_unit_stage_barrier() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    task_paths = [Path(path) for path in state["agent_task_paths"]]
    assert len(task_paths) == 2
    _bind_analysis_actions(state)

    first = read_json(task_paths[0])
    write_json(Path(first["result_path"]), _task_result(first))
    _validate_worker_task(task_paths[0])
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    second_task_id = progress["agent_sessions"]["analysis:U01"]["task_id"]
    state = resume_module_analysis(
        "smoke-01",
        str(data_root),
        settled_task_id=second_task_id,
    )
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert [action["unit_id"] for action in state["agent_actions"]] == ["U01"]
    assert all(read_json(path)["stage"] == "source_checkpoint" for path in task_paths)

    second = read_json(task_paths[1])
    write_json(Path(second["result_path"]), _task_result(second))
    _validate_worker_task(task_paths[1])
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"
    risk_task_paths = [Path(path) for path in state["agent_task_paths"]]
    assert set(task_paths).isdisjoint(risk_task_paths)
    assert all(read_json(path)["stage"] == "source_checkpoint" for path in task_paths)
    assert all(read_json(path)["stage"] == "risk_analysis" for path in risk_task_paths)
    assert {action["unit_id"] for action in state["agent_actions"]} == {"U00", "U01"}


def _graph_v2_limits_each_action_batch_to_eight() -> None:
    _, _, contract = _workspace(
        repositories=(
            "repo-a", "repo-b", "repo-c", "repo-d", "repo-e",
            "repo-f", "repo-g", "repo-h", "repo-i",
        )
    )
    data_root = Path(read_json(contract)["data_root"])
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert len(state["agent_actions"]) == 8
    assert len(state["agent_task_paths"]) == 8
    initial_actions = list(state["agent_actions"])

    untouched = run_module_analysis(str(contract))
    assert [action["unit_id"] for action in untouched["agent_actions"]] == [
        action["unit_id"] for action in initial_actions
    ]
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert progress["errors"] == []

    for index, action in enumerate(initial_actions):
        session_id = f"00000000-0000-4000-8000-{index + 1:012d}"
        subprocess.run(
            [
                sys.executable, "-m", "pangea_agent.cli.main", "record-agent-session",
                "--run-id", "smoke-01", "--data-root", str(data_root),
                "--role", "analysis", "--unit-id", action["unit_id"], "--task-id", session_id,
            ],
            check=True,
            capture_output=True,
        )

    for action in initial_actions[:3]:
        task_path = Path(action["task_path"])
        task = read_json(task_path)
        write_json(Path(task["result_path"]), _task_result(task))
        _validate_worker_task(task_path)

    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    assert [action["unit_id"] for action in state["agent_actions"]] == ["U08"]
    progress = read_json(data_root / "runs" / "smoke-01" / "progress.json")
    assert sum(
        session["status"] == "dispatched"
        for session in progress["agent_sessions"].values()
    ) == 5
    assert all(action["unit_id"] not in {item["unit_id"] for item in initial_actions[3:]} for action in state["agent_actions"])


def _graph_v2_cli_rejects_stale_stage_task() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _bind_analysis_actions(state)
    task_path = Path(state["agent_task_paths"][0])
    stale_source_task = read_json(task_path)
    write_json(Path(stale_source_task["result_path"]), _task_result(stale_source_task))
    _validate_worker_task(task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"
    risk_task_path = Path(state["agent_task_paths"][0])
    assert risk_task_path != task_path
    current_risk_task = read_json(risk_task_path)
    assert current_risk_task["stage"] == "risk_analysis"

    progress_path = data_root / "runs" / "smoke-01" / "progress.json"
    before_progress = progress_path.read_bytes()
    for command in ("prepare-worker-result", "validate-worker-result"):
        rejected = subprocess.run(
            [sys.executable, "-m", "pangea_agent.cli.main", command, "--task", str(task_path)],
            check=False,
            capture_output=True,
            text=True,
        )
        assert rejected.returncode != 0
        assert "不是当前 Graph 阶段任务" in rejected.stderr
        assert progress_path.read_bytes() == before_progress


def _graph_v2_legacy_resume_is_read_only() -> None:
    root = Path(tempfile.mkdtemp(prefix="pangea-legacy-resume-"))
    _TEMP_ROOTS.append(root)
    data_root = root / "data"
    complete_dir = data_root / "runs" / "legacy-complete"
    complete_dir.mkdir(parents=True)
    write_json(complete_dir / "progress.json", {
        "schema_version": "1.0",
        "run_id": "legacy-complete",
        "phase": "COMPLETE",
    })
    (complete_dir / "report.md").write_text("# legacy report\n", encoding="utf-8")
    (complete_dir / "report.html").write_text("<h1>legacy report</h1>\n", encoding="utf-8")
    before = {
        path.relative_to(complete_dir).as_posix(): path.read_bytes()
        for path in complete_dir.rglob("*")
        if path.is_file()
    }
    result = resume_module_analysis("legacy-complete", str(data_root))
    assert result["phase"] == "COMPLETE"
    assert result["report_path"] == str(complete_dir / "report.md")
    after = {
        path.relative_to(complete_dir).as_posix(): path.read_bytes()
        for path in complete_dir.rglob("*")
        if path.is_file()
    }
    assert after == before

    waiting_dir = data_root / "runs" / "legacy-waiting"
    waiting_dir.mkdir(parents=True)
    write_json(waiting_dir / "progress.json", {
        "schema_version": "1.0",
        "run_id": "legacy-waiting",
        "phase": "WAITING_ANALYSIS",
    })
    waiting_before = (waiting_dir / "progress.json").read_bytes()
    try:
        resume_module_analysis("legacy-waiting", str(data_root))
    except ValueError as exc:
        assert "旧版非终态不能继续" in str(exc)
    else:
        raise AssertionError("旧版非终态 Run 被 Graph V2 恢复")
    assert (waiting_dir / "progress.json").read_bytes() == waiting_before
    assert [path.name for path in waiting_dir.iterdir()] == ["progress.json"]


def _graph_v2_rejects_missing_or_future_completed_stage() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_SOURCE_CHECKPOINT"
    _bind_analysis_actions(state)
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    result_path = Path(task["result_path"])
    progress_path = data_root / "runs" / "smoke-01" / "progress.json"

    missing = _task_result(task)
    missing.pop("completed_stage")
    write_json(result_path, missing)
    assert "completed_stage" in _reject_worker_task(task_path)

    stage_spill = _task_result(task)
    evidence = {
        "chunk_id": "demo:module/entry.c:1-1",
        "location": "demo:module/entry.c:1-1",
        "observation": "不应在源码阶段提交的后续阶段证据",
    }
    stage_spill["evidence"] = [evidence]
    stage_spill["business_flows"] = [{
        "title": "提前生成的流程",
        "description": "源码阶段不应提交业务流程。",
        "steps": ["启动模块"],
        "mermaid": None,
        "evidence": [evidence],
    }]
    write_json(result_path, stage_spill)
    assert "源码 checkpoint 只能提交源码理解和 failure paths" in _reject_worker_task(task_path)

    future = _task_result(task)
    future["completed_stage"] = "risk_analysis"
    write_json(result_path, future)
    assert "当前 Graph 等待 source_checkpoint" in _reject_worker_task(task_path)

    write_json(result_path, _task_result(task))
    _validate_worker_task(task_path)
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_RISK_ANALYSIS"


def _cli_prints_custom_data_root() -> None:
    _, data_root, contract = _workspace()
    result = subprocess.run(
        [
            sys.executable,
            "-m",
            "pangea_agent.cli.main",
            "module-analysis",
            "--contract",
            str(contract),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    assert f"data_root={data_root}" in result.stdout


def _legacy_task_uses_c_cpp_defaults() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    payload = read_json(task_path)
    payload["unit"].pop("languages", None)
    payload["unit"].pop("frameworks", None)
    payload.pop("checkpoint_rubric_paths", None)
    write_json(task_path, payload)

    task = load_worker_task(task_path)
    assert task.unit.languages == ["c_cpp"]
    assert task.unit.frameworks == []
    assert task.checkpoint_rubric_paths == [
        "src/pangea_agent/rubrics/builtin/c_cpp_analysis.md"
    ]


def _lua_openubmc_task_metadata() -> None:
    _, _, contract = _lua_workspace()
    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))

    assert task["unit"]["languages"] == ["lua"]
    assert task["unit"]["frameworks"] == ["openubmc"]
    assert task["unit"]["source_scope"] == ["module/component.lua"]
    assert "module/helper.lua" in task["unit"]["context_scope"]
    assert task["checkpoint_rubric_paths"] == [
        "src/pangea_agent/rubrics/builtin/lua_analysis.md",
        "src/pangea_agent/rubrics/builtin/openubmc_analysis.md",
    ]
    check_ids = {item["check_id"] for item in task["semantic_check_items"]}
    assert any(value.startswith("SC-LUA-ERROR-") for value in check_ids)


def _semantic_check_missing_paths_reported_together() -> None:
    _, _, contract = _lua_workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    check_ids = [item["check_id"] for item in task["semantic_check_items"]]
    assert len(check_ids) >= 2
    result = _task_result(task)
    result["analysis_checkpoint"]["failure_paths"] = []
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(task_path),
        ],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "必须各有且只有一条" in rejected.stderr
    assert check_ids[0] in rejected.stderr and check_ids[1] in rejected.stderr
    assert any(value.startswith("SC-OPENUBMC-LIFECYCLE-") for value in check_ids)
    assert any(value.startswith("SC-OPENUBMC-SIGNAL-") for value in check_ids)
    assert any(value.startswith("SC-LUA-REQUIRE-") for value in check_ids)
    lifecycle_instruction = next(
        item["instruction"] for item in task["semantic_check_items"]
        if item["check_id"] == "SC-OPENUBMC-LIFECYCLE-01"
    )
    assert "注册行只改变 callback 表" in lifecycle_instruction
    assert "只有执行到 emit/dispatch" in lifecycle_instruction
    assert "config={} 缺少字段只会读到 nil" in lifecycle_instruction
    assert "只有 config=nil 等不可索引接收者" in lifecycle_instruction
    assert "主 check_id 只记录" in lifecycle_instruction
    assert "入口外层再包一层 pcall/xpcall" in lifecycle_instruction
    scenario_instructions = {
        item["check_id"]: item["instruction"]
        for item in task["semantic_check_items"]
    }
    assert "同实例初始化返工路径" in scenario_instructions[
        "SC-OPENUBMC-LIFECYCLE-01:retry"
    ]
    assert "config=nil 路径" in scenario_instructions[
        "SC-OPENUBMC-LIFECYCLE-01:nil-config"
    ]
    assert "未 start/init 就 update/emit" in scenario_instructions[
        "SC-OPENUBMC-LIFECYCLE-01:uninitialized-update"
    ]
    signal_instruction = next(
        item["instruction"] for item in task["semantic_check_items"]
        if item["check_id"].startswith("SC-OPENUBMC-SIGNAL-")
    )
    assert "绝对值为 2，不等于第二次调用的增量为 2" in signal_instruction
    assert "可重复调用的公开 create/ctor" in signal_instruction
    assert "双实例路径" in scenario_instructions[
        "SC-OPENUBMC-SIGNAL-01:multi-instance"
    ]
    error_instruction = next(
        item["instruction"] for item in task["semantic_check_items"]
        if item["check_id"].startswith("SC-LUA-ERROR-")
        and "失败前已到达的 callback 注册" in item["instruction"]
    )
    assert "connect/register 只把函数加入 callback 表" in error_instruction
    assert "函数体执行次数固定为 0" in error_instruction


def _lua_error_checks_separate_registration_from_callback_execution() -> None:
    unit = AnalysisUnit(
        unit_id="U00",
        repo_id="lua-demo",
        title="callback flow",
        source_scope=["module/app.lua"],
        context_scope=["module/component.lua"],
        focus=["callback"],
        dfx=["Reliability"],
        languages=["lua"],
        frameworks=["openubmc"],
    )
    inventory = {"files": [{
        "repo_id": "lua-demo",
        "path": "module/app.lua",
        "language": "lua",
        "calls": [
            {"line": 12, "callee": "instance:pre_init", "function_symbol": "app.start"},
            {"line": 13, "end_line": 15, "callee": "pcall", "function_symbol": "app.start"},
            {
                "line": 14,
                "callee": "instance:init",
                "function_symbol": "<anonymous@13>",
                "owner_function_symbol": "app.start",
            },
        ],
        "returns": [
            {
                "line": 16,
                "end_line": 16,
                "function_symbol": "app.start",
                "statement": "return ok, err",
            },
        ],
        "framework_signals": [],
    }, {
        "repo_id": "lua-demo",
        "path": "module/component.lua",
        "language": "lua",
        "calls": [
            {"line": 19, "callee": "self.changed:connect", "function_symbol": "Component:init"},
            {"line": 25, "callee": "error", "function_symbol": "Component:init"},
            {"line": 28, "callee": "self.changed:connect", "function_symbol": "Component:init"},
            {
                "line": 30,
                "callee": "error",
                "function_symbol": "<anonymous@28>",
                "owner_function_symbol": "Component:init",
            },
            {"line": 35, "callee": "self.changed:connect", "function_symbol": "Component:init"},
            {"line": 42, "callee": "self.changed:emit", "function_symbol": "Component:update"},
        ],
        "framework_signals": [
            {"kind": "class_declaration", "line": 3, "symbol": "Component"},
            {"kind": "class_lifecycle", "line": 6, "symbol": "Component:ctor"},
            {"kind": "class_lifecycle", "line": 14, "symbol": "Component:pre_init"},
            {"kind": "class_lifecycle", "line": 18, "symbol": "Component:init"},
            {"kind": "signal_callback", "line": 19, "symbol": "Component.changed"},
            {"kind": "signal_callback", "line": 28, "symbol": "Component.changed"},
            {"kind": "signal_callback", "line": 35, "symbol": "Component.changed"},
            {"kind": "signal_emit", "line": 42, "symbol": "Component.changed"},
        ],
    }]}
    checks = build_runtime_semantic_checks(unit, inventory)
    error_checks = [item for item in checks if item["check_id"].startswith("SC-LUA-ERROR-")]
    assert len(error_checks) == 2
    init_instruction = error_checks[0]["instruction"]
    assert "失败前已到达的 callback 注册：self.changed:connect@19" in init_instruction
    assert "失败后未到达的 callback 注册：self.changed:connect@28, self.changed:connect@35" in init_instruction
    callback_instruction = error_checks[1]["instruction"]
    assert "connect@28 注册的 callback 函数体内" in callback_instruction
    assert "Component.changed emit@42" in callback_instruction
    assert "不得把该错误写成 callback 注册阶段或 init 注册过程中的失败" in callback_instruction
    lifecycle_instruction = next(
        item["instruction"] for item in checks
        if item["check_id"] == "SC-OPENUBMC-LIFECYCLE-01"
    )
    assert "pcall@13 只保护其匿名函数内的 instance:init@14" in lifecycle_instruction
    assert "instance:pre_init@12 位于保护范围外" in lifecycle_instruction
    assert "错误中断前未完成的赋值保持原值" in lifecycle_instruction
    assert "公开入口直接调用时该错误会直接抛出" in lifecycle_instruction
    assert "app.start 原始语句 `return ok, err`" in lifecycle_instruction
    assert (
        "Component:init 注册顺序 [self.changed:connect@19, "
        "self.changed:connect@28, self.changed:connect@35]"
    ) in lifecycle_instruction
    assert "callback 内 self 仍指向注册它的实例，不会替换成 emit 发起实例" in lifecycle_instruction
    signal_instruction = next(
        item["instruction"] for item in checks
        if item["check_id"] == "SC-OPENUBMC-SIGNAL-01"
    )
    assert "每次 emit 都用本次 value 重新判断" in signal_instruction
    assert "A.C3 与全部 B callback 都未执行" in signal_instruction
    checks_by_id = {item["check_id"]: item["instruction"] for item in checks}
    assert "首次失败发生在 emit 前时" in checks_by_id[
        "SC-OPENUBMC-LIFECYCLE-01:retry"
    ]
    assert "公开入口直接调用会抛出错误" in checks_by_id[
        "SC-OPENUBMC-LIFECYCLE-01:nil-config"
    ]
    assert "现行资料没有承诺 nil 配置时" in checks_by_id[
        "SC-OPENUBMC-LIFECYCLE-01:nil-config"
    ]
    assert "排除为调用方误用" in checks_by_id[
        "SC-OPENUBMC-LIFECYCLE-01:uninitialized-update"
    ]
    multi_instruction = checks_by_id["SC-OPENUBMC-SIGNAL-01:multi-instance"]
    assert "第一次 normal 后 A/B" in multi_instruction
    assert "第二次 normal 后 A/B 都为 2" in multi_instruction
    assert "A 事件后 A=1/B=0" in multi_instruction


def _explicit_lua_scenarios_cannot_change_trigger() -> None:
    _, _, contract = _lua_workspace()
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    result = _task_result(task)
    retry_path = next(
        item for item in result["analysis_checkpoint"]["failure_paths"]
        if item["path_id"].endswith(":retry")
    )
    retry_path["trigger"] = "首次失败后重试成功"
    retry_path["side_effects"] = "完成重试"
    retry_path["final_states"] = "尚未发送事件"
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result", "--task", str(task_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "重试后单次 emit/update" in rejected.stderr

    result = _task_result(task)
    multi_path = next(
        item for item in result["analysis_checkpoint"]["failure_paths"]
        if item["path_id"].endswith(":multi-instance")
    )
    multi_path["trigger"] = "A emit trip 后 B emit normal"
    multi_path["side_effects"] = "先触发 trip，再触发 normal"
    multi_path["final_states"] = "记录两个实例状态"
    write_json(Path(task["result_path"]), result)
    rejected = subprocess.run(
        [sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result", "--task", str(task_path)],
        check=False,
        capture_output=True,
        text=True,
    )
    assert rejected.returncode != 0
    assert "A normal 后 B normal" in rejected.stderr


def _lua_direct_dependency_keeps_context_boundary_and_framework_checks() -> None:
    _, data_root, contract = _lua_workspace()
    repo = data_root / "repositories" / "lua-demo"
    (repo / "module" / "entry.lua").write_text(
        'return require("module.component")\n', encoding="utf-8"
    )
    (repo / "mc.lua").write_text(
        """local mc = {}
function mc.class() return {} end
function mc.signal()
    local signal = { callbacks = {} }
    function signal:emit(value)
        for _, callback in ipairs(self.callbacks) do
            local ok, err = pcall(callback, value)
            if not ok then return false, err end
        end
        return true, nil
    end
    return signal
end
return mc
""",
        encoding="utf-8",
    )
    subprocess.run(["git", "-C", str(repo), "add", "module/entry.lua", "mc.lua"], check=True)
    subprocess.run(
        [
            "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
            "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm", "lua entry fixture",
        ],
        check=True,
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.lua"]
    write_json(contract, payload)

    state = run_module_analysis(str(contract))
    task = read_json(Path(state["agent_task_paths"][0]))
    assert task["unit"]["source_scope"] == ["module/entry.lua"]
    assert "module/component.lua" in task["unit"]["context_scope"]
    assert "mc.lua" in task["unit"]["context_scope"]
    assert "module/helper.lua" not in task["unit"]["context_scope"]
    assert task["unit"]["frameworks"] == ["openubmc"]
    assert "src/pangea_agent/rubrics/builtin/openubmc_analysis.md" in task["checkpoint_rubric_paths"]
    component_checks = [
        item for item in task["semantic_check_items"]
        if item["subject_path"] == "module/component.lua"
    ]
    assert any(item["check_id"].startswith("SC-OPENUBMC-LIFECYCLE-") for item in component_checks)
    assert any(item["check_id"].startswith("SC-OPENUBMC-SIGNAL-") for item in component_checks)
    combined_instructions = "\n".join(item["instruction"] for item in component_checks)
    assert "条件 `not ok` 成立时执行原始语句 `return false, err`" in combined_instructions
    assert "后续循环项不再执行" in combined_instructions
    entry_inventory = next(
        item for item in state["inventory"]["files"]
        if item["repo_id"] == "lua-demo" and item["path"] == "module/entry.lua"
    )
    assert entry_inventory["imports"][0]["resolved_path"] == "module/component.lua"
    component_inventory = next(
        item for item in state["inventory"]["files"]
        if item["repo_id"] == "lua-demo" and item["path"] == "module/component.lua"
    )
    mc_import = next(
        item for item in component_inventory["imports"] if item["module"] == "mc"
    )
    assert mc_import["resolved_path"] == "mc.lua"
    markdown = render_report(state)
    assert "| 源码文件数 | 1 |" in markdown
    assert "| Lua 文件数 | 1 |" in markdown
    assert "覆盖 1 个源码文件（C/C++ 0，Lua 1）和 2 个上游语义文件" in markdown


def _explicit_lua_dependency_files_share_one_unit() -> None:
    _, _, contract = _lua_workspace()
    payload = read_json(contract)
    payload["source_scope"] = ["module/app.lua", "module/component.lua"]
    repo = Path(payload["data_root"]) / "repositories" / "lua-demo"
    (repo / "module" / "app.lua").write_text(
        'return require("module.component")\n', encoding="utf-8"
    )
    subprocess.run(["git", "-C", str(repo), "add", "module/app.lua"], check=True)
    subprocess.run(
        [
            "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
            "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm",
            "explicit lua dependency fixture",
        ],
        check=True,
    )
    write_json(contract, payload)

    state = run_module_analysis(str(contract))
    assert len(state["analysis_units"]) == 1
    unit = state["analysis_units"][0]
    assert unit["source_scope"] == ["module/app.lua", "module/component.lua"]
    assert "module/app.lua" not in unit["context_scope"]
    assert "module/component.lua" not in unit["context_scope"]


def _lua_parser_binds_direct_assignments_and_self_signals() -> None:
    root = Path(tempfile.mkdtemp(prefix="pangea-lua-symbols-"))
    _TEMP_ROOTS.append(root)
    path = root / "component.lua"
    path.write_text(
        """local ignored, Component = make(), mc.class("Component")
local Wrapped = decorate(mc.class("Wrapped"))
Component.changed = mc.signal()
function Component:init() self.changed:connect(function() end) end
function Component:run() self.changed:emit() end
""",
        encoding="utf-8",
    )
    parsed = parse_lua_file(path)
    declarations = [
        item["symbol"] for item in parsed["framework_signals"]
        if item["kind"] == "class_declaration"
    ]
    assert "Component" in declarations
    assert "ignored" not in declarations and "Wrapped" not in declarations
    signal_kinds = {item["kind"] for item in parsed["framework_signals"]}
    assert {"signal_callback", "signal_emit"} <= signal_kinds

    scoped_path = root / "scoped.lua"
    scoped_path.write_text(
        """function A()
    local changed = mc.signal()
    changed:emit()
end
function B()
    changed:emit()
end
function C:init()
    self.changed:connect(function()
        error("callback failed")
    end)
end
local M = { run = function() return true end }
return M
""",
        encoding="utf-8",
    )
    scoped = parse_lua_file(scoped_path)
    emits = [item for item in scoped["framework_signals"] if item["kind"] == "signal_emit"]
    assert len(emits) == 1
    assert any(item["symbol"] == "M.run" for item in scoped["functions"])
    assert any(
        item.get("function_symbol") == "M.run"
        and item["statement"] == "return true"
        for item in scoped["returns"]
    )
    nested_error = next(item for item in scoped["calls"] if item["callee"] == "error")
    assert nested_error["function_symbol"].startswith("<anonymous@")
    assert nested_error["owner_function_symbol"] == "C:init"

    signal_path = root / "signal.lua"
    signal_path.write_text(
        """function signal:emit()
    for _, callback in ipairs(self.callbacks) do
        local ok, err = pcall(callback)
        if not ok then return false, err end
    end
    return true, nil
end
""",
        encoding="utf-8",
    )
    signal_returns = parse_lua_file(signal_path)["returns"]
    guarded = next(item for item in signal_returns if item["statement"] == "return false, err")
    assert guarded["guard"] == "not ok"


def _lua_coverage_path_disambiguates_duplicate_symbols() -> None:
    inventory = {
        "files": [
            {
                "repo_id": "lua-demo",
                "path": "module/a.lua",
                "language": "lua",
                "functions": [{"symbol": "Alpha:init", "line": 2}],
            },
            {
                "repo_id": "lua-demo",
                "path": "module/b.lua",
                "language": "lua",
                "functions": [{"symbol": "Beta:init", "line": 4}],
            },
        ]
    }
    base = {"coverage_type": "function", "module": "", "function": "init", "count": 0}
    without_path = match_coverage_records([base], inventory)
    assert len(without_path["unmatched"]) == 1

    matched = match_coverage_records([{**base, "path": "module/b.lua"}], inventory)
    assert matched["matched"][0]["matches"] == [
        {"repo_id": "lua-demo", "path": "module/b.lua", "line": 4}
    ]

    unmatched = match_coverage_records([{**base, "path": "module/missing.lua"}], inventory)
    assert len(unmatched["unmatched"]) == 1
    assert not unmatched["matched"]
    dotted = match_coverage_records([
        {**base, "function": "Beta.init"}
    ], inventory)
    assert dotted["matched"][0]["matches"][0]["line"] == 4


def _legacy_inventory_coverage_defaults_to_c_cpp() -> None:
    files = [{
        "repo_id": "repo-a",
        "path": "src/legacy.c",
        "functions": [{"symbol": "legacy_start", "line": 12}],
    }]
    record = {
        "coverage_type": "function",
        "module": "",
        "path": "",
        "function": "legacy_start",
        "count": 0,
    }
    for legacy_file in (files[0], {**files[0], "language": "c_cpp"}):
        report = match_coverage_records([record], {"files": [legacy_file]})
        assert len(report["matched"]) == 1
        assert not report["ambiguous"]
        assert not report["unmatched"]


def _coverage_ignores_context_symbol_collisions() -> None:
    unit = AnalysisUnit(
        unit_id="U00",
        repo_id="lua-demo",
        title="Lua source",
        source_scope=["module/source.lua"],
        context_scope=["module/context.lua"],
        focus=["test_cases"],
        dfx=["功能与状态"],
        languages=["lua"],
    )
    inventory = {
        "files": [
            {
                "repo_id": "lua-demo",
                "path": "module/source.lua",
                "language": "lua",
                "functions": [{"symbol": "Component:init", "line": 3}],
            },
            {
                "repo_id": "lua-demo",
                "path": "module/context.lua",
                "language": "lua",
                "functions": [{"symbol": "Component:init", "line": 8}],
            },
        ],
        "file_count": 2,
    }
    records = [{
        "coverage_type": "function",
        "module": "component",
        "function": "Component.init",
        "count": 0,
    }]
    report = match_coverage_records(
        records,
        _source_inventory(unit, inventory),
        path_inventory=inventory,
    )
    context = _coverage_context(unit, report)
    assert len(report["matched"]) == 1 and not report["ambiguous"]
    assert context[0]["path"] == "module/source.lua"

    exact_context = match_coverage_records(
        [{**records[0], "path": "module/context.lua"}],
        _source_inventory(unit, inventory),
        path_inventory=inventory,
    )
    exact_context_items = _coverage_context(unit, exact_context)
    assert exact_context_items == []


def _lua_context_path_coverage_reaches_worker() -> None:
    _, data_root, contract = _lua_workspace()
    repo = data_root / "repositories" / "lua-demo"
    (repo / "module" / "entry.lua").write_text(
        'return require("module.component")\n',
        encoding="utf-8",
    )
    subprocess.run(["git", "-C", str(repo), "add", "module/entry.lua"], check=True)
    subprocess.run(
        [
            "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
            "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm",
            "context coverage fixture",
        ],
        check=True,
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.lua"]
    write_json(contract, payload)
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["module", "path", "function", "count"])
    sheet.append(["component", "module/component.lua", "Component:run", 0])
    workbook.save(coverage / "context.xlsx")
    workbook.close()

    state = run_module_analysis(str(contract))
    assert state["coverage_report"]["matched"][0]["matches"][0]["path"] == "module/component.lua"
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    assert task["coverage_context"] == []


def _source_prefix_evidence_is_normalized() -> None:
    _, data_root, contract = _workspace()
    state = run_module_analysis(str(contract))
    _bind_analysis_actions(state)
    source_task_path = Path(state["agent_task_paths"][0])
    source_task = read_json(source_task_path)
    write_json(Path(source_task["result_path"]), _task_result(source_task))
    _validate_worker_task(source_task_path)
    state = run_module_analysis(str(contract))
    task_path = Path(state["agent_task_paths"][0])
    task = read_json(task_path)
    result = _task_result(task)
    result["evidence"][0]["chunk_id"] = "repo:demo:module/entry.c:1-1"
    result["evidence"][0]["location"] = None
    result["business_flows"][0]["evidence"][0]["chunk_id"] = "source:module/entry.c:1-1"
    result["business_flows"][0]["evidence"][0]["location"] = None
    write_json(Path(task["result_path"]), result)
    subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "validate-worker-result",
            "--task", str(task_path),
        ],
        check=True,
        capture_output=True,
    )
    normalized = read_json(Path(task["result_path"]))
    assert normalized["evidence"][0]["chunk_id"].startswith("demo:module/entry.c:")
    assert normalized["evidence"][0]["status"] == "confirmed"
    assert normalized["business_flows"][0]["evidence"][0]["status"] == "confirmed"


def _coverage_context_collision_keeps_source_gap_through_advance() -> None:
    _, data_root, contract = _lua_workspace()
    repo = data_root / "repositories" / "lua-demo"
    (repo / "module" / "entry.lua").write_text(
        """local Context = require("module.component")
local Component = {}
function Component:init() return Context ~= nil end
return Component
""",
        encoding="utf-8",
    )
    subprocess.run(["git", "-C", str(repo), "add", "module/entry.lua"], check=True)
    subprocess.run(
        [
            "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
            "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm", "coverage collision fixture",
        ],
        check=True,
    )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.lua"]
    write_json(contract, payload)
    coverage = data_root / "coverage"
    coverage.mkdir()
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["module", "function", "count"])
    sheet.append(["component", "Component.init", 0])
    workbook.save(coverage / "collision.xlsx")
    workbook.close()

    state = run_module_analysis(str(contract))
    assert len(state["coverage_report"]["matched"]) == 1
    assert not state["coverage_report"]["ambiguous"]
    state = _advance_to_test_generation(state)
    task = read_json(Path(state["agent_task_paths"][0]))
    gap_id = task["coverage_context"][0]["gaps"][0]["coverage_id"]
    assert task["coverage_context"][0]["path"] == "module/entry.lua"
    result = _task_result(task)
    result["test_cases"] = [_test_case("U00-COV-TC1", coverage_ids=[gap_id])]
    result["analysis_checkpoint"]["coverage_decisions"] = [{
        "coverage_id": gap_id,
        "disposition": "new_coverage_case",
        "linked_test_case_ids": ["U00-COV-TC1"],
        "reason": "从当前源码入口补齐未执行函数。",
    }]
    write_json(Path(task["result_path"]), result)
    _validate_worker_task(Path(state["agent_task_paths"][0]))
    state = run_module_analysis(str(contract))
    assert state["phase"] == "WAITING_INDEPENDENT_REVIEW"
    assert len(state["coverage_report"]["matched"]) == 1
    assert not state["coverage_report"]["ambiguous"]


def _lua_context_inventory_isolated_by_repository() -> None:
    _, data_root, contract = _workspace(repositories=("repo-a", "repo-b"))
    for repo_id in ("repo-a", "repo-b"):
        repo = data_root / "repositories" / repo_id
        (repo / "module" / "entry.c").unlink()
        (repo / "module" / "component.lua").write_text(
            "local C = mc.class(\"C\")\nfunction C:init() end\nreturn C\n",
            encoding="utf-8",
        )
        entry = (
            'return require("module.component")\n'
            if repo_id == "repo-a"
            else "return { value = true }\n"
        )
        (repo / "module" / "entry.lua").write_text(entry, encoding="utf-8")
        subprocess.run(["git", "-C", str(repo), "add", "-A"], check=True)
        subprocess.run(
            [
                "git", "-C", str(repo), "-c", "user.name=PANGEA Smoke",
                "-c", "user.email=pangea-smoke@example.invalid", "commit", "-qm", "lua repo fixture",
            ],
            check=True,
        )
    payload = read_json(contract)
    payload["source_scope"] = ["module/entry.lua"]
    write_json(contract, payload)
    state = run_module_analysis(str(contract))
    inventory_paths = {
        (item["repo_id"], item["path"]) for item in state["inventory"]["files"]
    }
    assert ("repo-a", "module/component.lua") in inventory_paths
    assert ("repo-b", "module/component.lua") not in inventory_paths


def _coverage_workbook_preserves_source_path() -> None:
    root = Path(tempfile.mkdtemp(prefix="pangea-coverage-path-"))
    _TEMP_ROOTS.append(root)
    workbook = Workbook()
    function_sheet = workbook.active
    function_sheet.title = "functions"
    function_sheet.append(["module", "path", "function", "count"])
    function_sheet.append(["power", "module/component.lua", "init", 0])
    branch_sheet = workbook.create_sheet("branches")
    branch_sheet.append([
        "branch_id", "路径", "function", "condition", "true_count", "false_count"
    ])
    branch_sheet.append(["B1", "module/component.lua", "run", "ready", 0, 2])
    path = root / "coverage.xlsx"
    workbook.save(path)
    workbook.close()

    records, warnings = parse_coverage_xlsx(path)
    assert warnings == []
    assert [item["path"] for item in records] == [
        "module/component.lua", "module/component.lua"
    ]


def _mixed_language_reports_show_frameworks() -> None:
    _, _, contract = _workspace()
    state = run_module_analysis(str(contract))
    state["inventory"]["files"].append({
        "repo_id": "demo",
        "path": "module/component.lua",
        "language": "lua",
        "frameworks": ["openubmc"],
    })
    state["inventory"]["file_count"] += 1
    state["analysis_units"][0]["source_scope"].append("module/component.lua")
    state["analysis_units"][0]["languages"] = ["c_cpp", "lua"]
    state["analysis_units"][0]["frameworks"] = ["openubmc"]

    markdown = render_report(state)
    html = render_html_report(state)
    assert "| C/C++ 文件数 | 1 |" in markdown
    assert "| Lua 文件数 | 1 |" in markdown
    assert "| 识别框架 | openubmc |" in markdown
    assert "<td>C/C++ 文件数</td><td>1</td>" in html
    assert "<td>Lua 文件数</td><td>1</td>" in html
    assert "<td>识别框架</td><td>openubmc</td>" in html

    legacy_state = run_module_analysis(str(contract))
    for item in legacy_state["inventory"]["files"]:
        item.pop("language", None)
    legacy_markdown = render_report(legacy_state)
    assert "| C/C++ 文件数 | 1 |" in legacy_markdown


def _lua_callback_methodology_reaches_all_clients() -> None:
    root = Path(__file__).resolve().parents[1]
    lua_rubric = (root / "src/pangea_agent/rubrics/builtin/lua_analysis.md").read_text()
    framework_rubric = (
        root / "src/pangea_agent/rubrics/builtin/openubmc_analysis.md"
    ).read_text()
    test_rubric = (
        root / "src/pangea_agent/rubrics/builtin/test_case_generation.md"
    ).read_text()
    assert "没有发生 emit/dispatch 时" in lua_rubric
    assert "`callback_count` 只增加 2，不是 4" in lua_rubric
    assert "`repo_id:path:line + 错误点` 为键横向核对所有 failure path" in lua_rubric
    assert "expected result 才能写外层 wrapper 的 `ok=false`" in lua_rubric
    assert "兄弟场景成立不能把" in lua_rubric
    assert "增量仍是 1" in lua_rubric
    assert "A 注册的 callback 写 A，B 注册的 callback" in lua_rubric
    assert "不能把上一次的分支结果当成 callback 的永久状态" in lua_rubric
    assert "`normal` emit 不能沿用上一次错误结论" in framework_rubric
    assert "A 为 2、1、true" in lua_rubric and "B 为 1、1、true" in lua_rubric
    assert "TestCase 的通过预期仍应写 `callback_count` 增量 1" in framework_rubric
    assert "类表或模块表共享 signal 是实现事实，不自动等于产品风险" in framework_rubric
    assert "create/ctor 入口" in framework_rubric
    assert '`steps=[{"action":"业务动作","expected_result":"正确产品结果","failure_observation":"当前错误观测，出现即 FAIL"}]`' in test_rubric
    assert "不得把动作和预期拆到两个数组" in test_rubric
    assert "expected_results" not in test_rubric

    for relative_path in (
        ".opencode/agents/analysis-worker.md",
        ".claude/agents/analysis-worker.md",
    ):
        rules = (root / relative_path).read_text()
        assert "表长度为 4" in rules and "实际增量为 2" in rules
        assert "TestCase" in rules and "正确" in rules and "增量 1" in rules
        assert '`steps=[{"action":"业务动作","expected_result":"正确产品结果","failure_observation":"当前错误观测，出现即 FAIL"}]`' in rules
        assert "expected_results" not in rules
        assert "记录/读取/验证本次返回值" in rules
        assert "不是只查 issue 点名的 `path_id`" in rules
        assert "`repo_id:path:line + 错误点` 为键做一次横向一致性检查" in rules
        assert "`prior_result_path` 只读" in rules
        assert "整个公开入口外再包一层 pcall/xpcall" in rules or "整个入口外再包一层" in rules
        assert "第二次后绝对值为 2" in rules
        assert "review issue 是待核对的修复请求，不是源码或资料证据" in rules
        assert "不得为了迎合 reviewer 反转执行方向" in rules
        assert "所有 edit/write 的目标都必须是该" in rules
        assert "A 为 2/1/true、B 为 1/1/true" in rules
    for relative_path in (
        ".opencode/agents/review-worker.md",
        ".claude/agents/review-worker.md",
    ):
        rules = (root / relative_path).read_text()
        assert "尚未 emit/dispatch" in rules
        assert "数字类 `prior_issues.required_change` 不是验证 oracle" in rules
        assert "TestCase 的正确预期是增量 1" in rules
        assert "记录/读取/验证本次返回值" in rules
        assert "不能先放行到返工验证才发现第二条" in rules
        assert "公开 create/ctor" in rules
        assert "第二次后绝对值为 2" in rules
        assert "句中任何位置出现“当前实现”" in rules
        assert "只能要求修改 worker result" in rules
        assert "两边同为“不执行”" in rules
        assert "某 callback 在 `trip` 分支失败" in rules
        assert "不得发 issue 要求把 B=2/2 写成" in rules


def _graph_control_contract_reaches_all_clients() -> None:
    root = Path(__file__).resolve().parents[1]
    controller_paths = (
        "AGENTS.md",
        "CLAUDE.md",
        ".agents/skills/pangea-agent/SKILL.md",
        ".agents/pangea/dsh.md",
        ".opencode/agents/pangea-agent.md",
        ".opencode/commands/module-analysis.md",
        ".claude/commands/module_analysis.md",
    )
    combined = "\n".join((root / path).read_text() for path in controller_paths)
    assert "after_completion=resume_run" in combined
    for forbidden in (
        "record_completion_then_resume",
        "--status completed",
        "--status dispatched",
        "list_agents",
        "job_output",
        "[STAGE:",
    ):
        assert forbidden not in combined
    for path in controller_paths:
        rules = (root / path).read_text()
        assert "action" in rules
    dsh_skill = (root / ".agents/skills/pangea-agent/SKILL.md").read_text()
    assert (
        "--role analysis --unit-id <unit_id> --task-id <subagent_id>"
        in dsh_skill
    )
    assert "After the command succeeds, end the DSH root turn immediately" in dsh_skill
    assert "A `subagent-report` is informational only" in dsh_skill
    assert "A `subagent-settled` notice for the current action is only a wake-up signal" in dsh_skill
    assert "--settled-task-id <the subagent_id bound to this action>" in dsh_skill
    assert "without inferring PASS" in dsh_skill
    assert "never compare actions to decide whether to stop" in dsh_skill
    assert "the same action, stop truthfully" not in dsh_skill
    dsh_adapter = (root / ".agents/pangea/dsh.md").read_text()
    assert "统一插件 `dsh-pangea` 内置的唤醒策略" in dsh_adapter
    assert "成功返回的 action 一律按其内容执行" in dsh_adapter
    assert "--settled-task-id <本次 action 绑定的 subagent_id>" in dsh_adapter
    dsh_analysis_rules = (root / ".opencode/agents/analysis-worker.md").read_text()
    assert "禁止把 `0 >= 0`" in dsh_analysis_rules
    assert "返回 `void` 的 API 不能虚构失败返回" in dsh_analysis_rules
    assert "callee 直接 `return` 不等于 caller 跳过清理" in dsh_analysis_rules
    assert "返回值只为兼容且固定为 0" in dsh_analysis_rules
    assert "`for (++ptr; ptr < end; ptr++)` 又清零全部后续元素" in dsh_analysis_rules
    assert "正确产品在同一次故障注入下的通过标准" in dsh_analysis_rules
    assert "失败时最后一条赋值不可达" in dsh_analysis_rules
    assert "不得仅凭“应该允许”建立兼容风险" in dsh_analysis_rules
    dsh_review_rules = (root / ".opencode/agents/review-worker.md").read_text()
    assert "`0 >= 0` 必须判 true" in dsh_review_rules
    assert "返回 `void` 的 API 不存在可检查的失败返回" in dsh_review_rules
    assert "callee 直接 `return` 不能作为 caller 跳过清理" in dsh_review_rules
    assert "`/proc/iomem` 只表示物理资源登记" in dsh_review_rules
    assert "返回值仅兼容且固定为 0" in dsh_review_rules
    assert "`++ptr` 正是移动到第一个后续元素" in dsh_review_rules
    assert "必须使用完全相同的 trigger" in dsh_review_rules
    assert "NULL 分支不会执行最后赋值" in dsh_review_rules
    assert "不得要求修改被测源码来解决分析报告" in dsh_review_rules
    for relative_path in (
        ".opencode/agents/analysis-worker.md",
        ".claude/agents/analysis-worker.md",
    ):
        rules = (root / relative_path).read_text()
        assert "“使用新实例重新测试”不能作为清理" in rules
        assert "稳定消息片段" in rules
        assert "不得写 `err='<完整字符串>'`" in rules
        assert "禁止使用 Write 整体覆盖" in rules
        assert "禁止用 Bash、Python、正则或临时脚本批量重写/修复 JSON" in rules
        assert "触发风险的调用步骤必须同时写完整返回值和完整调用后状态" in rules
        assert "普通“失败后修复重试”只执行" in rules
        assert "两个局部变量在所有能到达使用点的路径上都由同一次赋值保持相等" in rules
        assert "Graph 要求重新提交当前 task" in rules
    for relative_path in (
        ".opencode/agents/review-worker.md",
        ".claude/agents/review-worker.md",
    ):
        rules = (root / relative_path).read_text()
        assert "在填写任何 `covered` 前先逐条检查" in rules
        assert "正确实现满足这条 expected 时是否 PASS" in rules
        assert "可以使用空关联数组" in rules
        assert "它们只是上轮结论，不是本轮答案" in rules
        assert "禁止使用 Write 整体覆盖" in rules
        assert "禁止用 Bash、Python、正则或临时脚本批量重写/修复 JSON" in rules
        assert "后续“检查状态”中的任何字段与触发调用真实终态" in rules
        assert "当前实现必须至少违反" in rules
        assert "相同调用连续出现几次就累计几次注册/计数" in rules
        assert "两个局部变量若在所有到达使用点的路径上由同一次赋值保持相等" in rules
        assert "Graph 要求重新提交当前 task" in rules
    lua_rules = (root / "src/pangea_agent/rubrics/builtin/lua_analysis.md").read_text()
    assert "该函数体同步调用的全部下层函数" in lua_rules
    assert "只要它是在受保护函数的动态调用链中执行，就会被捕获" in lua_rules
    assert "连续 `n` 次失败" in lua_rules
    assert "完成状态只能由当前 task 的提交校验记录" in (
        root / "src/pangea_agent/graph/nodes/apply_run_event.py"
    ).read_text()


def _companion_adapter_advances_current_graph() -> None:
    _, data_root, contract = _workspace()
    created_process = subprocess.run(
        [
            sys.executable, "-m", "pangea_agent.cli.main", "runs", "create",
            "--contract", str(contract),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    envelope = json.loads(created_process.stdout.strip().splitlines()[-1])
    assert envelope["api_version"] == "1.0" and envelope["ok"] is True
    created = envelope["result"]
    action = created["agent_actions"][0]
    assert action["action_id"].endswith("|analysis:U00|source_checkpoint")
    bound = bind_action(
        str(data_root), created["run_id"], action["action_id"], _ANALYSIS_SESSION_ID
    )
    assert bound["status"] == "dispatched"
    invalid = validate_action(str(data_root), created["run_id"], action["action_id"])
    assert invalid["status"] == "invalid"
    task_path = Path(action["task_path"])
    task = read_json(task_path)
    write_json(Path(task["result_path"]), _task_result(task))
    _validate_worker_task(task_path)
    assert validate_action(
        str(data_root), created["run_id"], action["action_id"]
    )["status"] == "valid"
    advanced = settle_action(str(data_root), created["run_id"], action["action_id"])
    assert advanced["phase"] == "WAITING_RISK_ANALYSIS"
    assert advanced["agent_actions"][0]["action"] == "continue_agent"
    assert advanced["agent_actions"][0]["action_id"].endswith(
        "|analysis:U00|risk_analysis"
    )


SCENARIOS: tuple[tuple[str, Scenario], ...] = (
    ("PASS 到双报告", _pass_report),
    ("独立复核后替换失效 Worker 且不重复审计", _late_worker_rejection_reuses_independent_review),
    ("reviewer 发现遗漏时不能 PASS", _review_missing_finding_cannot_pass),
    ("阻塞 finding 必须进入返工 issue", _review_blocking_finding_must_reach_rework_issue),
    ("对照复核必须逐项绑定 worker 风险与用例", _comparison_must_account_for_every_worker_artifact),
    ("对照复核必须逐条记录 TestCase 通过标准", _comparison_requires_test_case_oracle_checks),
    ("空 Coverage 输入时 review 不得伪造闭环", _comparison_cannot_invent_coverage_closure),
    ("review issue 不得把错误实现改成测试通过标准", _comparison_cannot_replace_oracle_with_buggy_behavior),
    ("reviewer 自身误判不得制造 worker 返工", _comparison_does_not_rework_worker_for_reviewer_mistake),
    ("对照复核不能丢弃独立 finding", _comparison_cannot_drop_independent_findings),
    ("REWORK 同 reviewer 通过", _rework_same_reviewer),
    ("返工任务保留冻结资料路径", _rework_preserves_allowed_material_paths),
    ("无效返工直接终止而不循环 reviewer", _invalid_rework_finishes_unresolved_without_reviewer_loop),
    ("原 reviewer 无法恢复时显式 UNRESOLVED", _reviewer_unavailable_has_explicit_unresolved_path),
    ("截断结果覆盖修正", _truncated_correction),
    ("黑盒步骤必须直接绑定预期", _unpaired_test_step_rejected),
    ("黑盒步骤与预期成对渲染", _paired_test_steps_render_together),
    ("黑盒预期不得混入当前错误实现", _expected_result_cannot_include_current_behavior),
    ("需求补测无需伪造风险关联", _requirement_only_test_case_is_accepted),
    ("semantic check 约束风险实现范围", _semantic_check_risk_scope_is_enforced),
    ("派生 semantic path 不改写主路径", _derived_semantic_path_does_not_rewrite_main_path),
    ("无法关联证据不能离开 worker 阶段", _unmatched_evidence_is_pending),
    ("未确认风险证据不能离开 worker 阶段", _pass_rejects_unconfirmed_risk_evidence),
    ("source 前缀证据规范化为真实源码引用", _source_prefix_evidence_is_normalized),
    ("跨单元重复 ID 自动修正", _duplicate_ids_correction),
    ("机械路径变化自动修正", _mechanical_task_change_does_not_block),
    ("review 机械字段自动修正", _mechanical_review_fields_do_not_block),
    ("不存在 scope 拒绝", _missing_scope),
    ("终态报告恢复", _report_recovery),
    ("文档缺口强制不完整", _document_gap),
    ("有 Coverage 记录但无缺口时不强制补测", _coverage_reference_only),
    ("Coverage 缺口必须逐项闭环且用例关联严格", _coverage_gap_requires_closure_and_linkage_is_strict),
    ("Coverage 报告只展示当前范围函数级零覆盖", _coverage_report_only_shows_current_zero_functions),
    ("空 Coverage 输入拒绝伪造缺口结论", _empty_coverage_rejects_claimed_gap),
    ("风险阶段可先规划 Coverage 且用例阶段必须闭环", _risk_stage_can_plan_coverage_before_test_ids_exist),
    ("相关资料必须闭环且报告不展示排除章节", _material_traceability_report),
    ("已确认历史问题作为只读回归参考进入 Run", _confirmed_historical_issues_are_frozen_as_references),
    ("多 repo 单元隔离", _multi_repo_isolation),
    ("返工期间未返工结果编辑不阻塞", _unchanged_result_edit_does_not_block),
    ("Run ID 并发原子占位", _run_ids_are_reserved_atomically),
    ("module-analysis 自动清理本会话 pending contract", _module_analysis_owns_pending_contract_cleanup),
    ("Agent 路径统一且非法 source_scope 拒绝", _agent_paths_and_contract_scope_are_posix),
    ("目录 scope 按文件族拆分并补充唯一直接被调函数", _directory_scope_splits_file_families_and_adds_unique_direct_callee),
    ("超过目标规模的不同职责文件不合并", _large_distinct_families_do_not_merge_into_one_unit),
    ("直接调用者上下文数量有界", _direct_caller_context_is_bounded),
    ("预处理宏不会吞并后续函数定义", _preprocessor_macro_is_not_a_function_definition),
    ("多行条件不会冒充函数定义", _multiline_condition_is_not_a_function_definition),
    ("块宏不会冒充导出函数", _block_macro_is_not_an_exported_symbol),
    ("超大直接依赖不冻结整文件", _oversized_direct_callee_is_not_frozen),
    ("专项方法只按当前源码范围选择", _specialized_rubrics_follow_source_scope),
    ("DOCA submit 失败清理进入语义检查", _doca_submit_cleanup_becomes_a_semantic_check),
    ("显式 scope 外实现只作为上下文", _explicit_scope_keeps_external_implementation_as_context),
    ("范围只扩到直接调用与相关上下文", _bounded_scope_expansion),
    ("状态上下文均衡保留生命周期与重配置", _state_context_balances_lifecycle_and_reconfiguration),
    ("已选头文件保留直接本地依赖", _selected_header_keeps_its_direct_local_dependency),
    ("预期行为不能列为风险", _expected_behavior_not_risk),
    ("普通源码目录不因缺少 Git 版本阻塞交付", _unversioned_source_is_deliverable),
    ("SOURCE_READY 恢复只使用冻结输入", _source_checkpoint_uses_frozen_inputs),
    ("源码 checkpoint 提前拒绝已知 AE4DMA 误读", _source_checkpoint_rejects_known_ae4dma_misreads),
    ("INDEX_READY 恢复不再读取活动源码", _index_checkpoint_resumes_without_live_source),
    ("Agent 启动状态写入 Run checkpoint", _agent_start_checkpoint),
    ("Agent 会话拒绝伪造 task_id", _record_agent_session_rejects_fake_task_id),
    ("并发 Worker 校验不会丢失完成状态", _parallel_worker_validation_preserves_completions),
    ("快速 Worker 先完成后绑定仍可推进", _fast_worker_completion_survives_late_session_binding),
    ("未通过提交校验的 Worker 由 Graph 返回续接 action", _unvalidated_worker_returns_graph_continue_action),
    ("未执行 callback 不得被归因写入状态", _nonexecuted_callback_cannot_write_state),
    ("Graph V2 阶段与 Agent action 严格推进到 COMPLETE", _graph_v2_stage_actions_reach_complete),
    ("Graph V2 两单元阶段 barrier", _graph_v2_two_unit_stage_barrier),
    ("Graph V2 每批最多返回八个 Agent action", _graph_v2_limits_each_action_batch_to_eight),
    ("Graph V2 CLI 拒绝旧阶段 task", _graph_v2_cli_rejects_stale_stage_task),
    ("Graph V2 只读旧 COMPLETE 并拒绝旧非终态", _graph_v2_legacy_resume_is_read_only),
    ("Graph V2 拒绝缺失或越级 completed_stage", _graph_v2_rejects_missing_or_future_completed_stage),
    ("CLI 回显自定义 data_root", _cli_prints_custom_data_root),
    ("已知 C 宏解析误报不冒充真实缺口", _known_c_macro_parse_artifacts),
    ("返工复核骨架移除已删除产物关联", _rework_review_prunes_removed_links),
    ("复核问题必须给出可验证的确定修改", _review_issue_requires_actionable_change),
    ("字面边界条件不得写反", _literal_boundary_comparison_cannot_be_reversed),
    ("旧 WorkerTask 使用 C/C++ 默认规则", _legacy_task_uses_c_cpp_defaults),
    ("Lua openUBMC task 冻结语言与规则", _lua_openubmc_task_metadata),
    ("semantic check 缺失路径一次性报全", _semantic_check_missing_paths_reported_together),
    ("Lua error 检查区分注册阶段与 callback 执行阶段", _lua_error_checks_separate_registration_from_callback_execution),
    ("Lua 显式场景不得改换触发序列", _explicit_lua_scenarios_cannot_change_trigger),
    ("Lua 直接依赖保留 context 边界并生成框架检查", _lua_direct_dependency_keeps_context_boundary_and_framework_checks),
    ("显式 Lua 依赖文件合并为同一分析单元", _explicit_lua_dependency_files_share_one_unit),
    ("Lua parser 正确绑定多赋值与 self signal", _lua_parser_binds_direct_assignments_and_self_signals),
    ("Lua 重名函数 Coverage 使用路径消歧", _lua_coverage_path_disambiguates_duplicate_symbols),
    ("旧 inventory Coverage 默认按 C/C++ 匹配", _legacy_inventory_coverage_defaults_to_c_cpp),
    ("Coverage 匹配忽略 context 重名符号", _coverage_ignores_context_symbol_collisions),
    ("带路径 Lua context Coverage 进入 worker", _lua_context_path_coverage_reaches_worker),
    ("Coverage source 缺口推进后仍保持匹配", _coverage_context_collision_keeps_source_gap_through_advance),
    ("Lua context inventory 按源码仓隔离", _lua_context_inventory_isolated_by_repository),
    ("Coverage Excel 保留源码路径", _coverage_workbook_preserves_source_path),
    ("混合语言报告显示框架", _mixed_language_reports_show_frameworks),
    ("Lua callback 执行账本同步到各客户端", _lua_callback_methodology_reaches_all_clients),
    ("Graph 控制契约同步到各客户端", _graph_control_contract_reaches_all_clients),
    ("Companion adapter 推进当前 Graph", _companion_adapter_advances_current_graph),
)


def main() -> None:
    logging.getLogger("pypdf").setLevel(logging.CRITICAL)
    for name, scenario in SCENARIOS:
        scenario()
        print(f"PASS  {name}")
    print(f"PASS  Graph V2 smoke ({len(SCENARIOS)} scenarios)")


if __name__ == "__main__":
    main()
