from __future__ import annotations

import mimetypes
import re
import shutil
import zipfile
from pathlib import Path

from .types import DocumentExtraction, EvidenceAttachment


class DependencyUnavailableError(RuntimeError):
    def __init__(self, package: str, document_type: str) -> None:
        self.package = package
        self.document_type = document_type
        super().__init__(f"{document_type} extraction requires the '{package}' package")


def _attachment_dir(source: Path, attachments_root: Path) -> Path:
    materials_root = attachments_root.parent / "frozen" / "materials"
    try:
        relative = source.resolve().relative_to(materials_root.resolve())
    except ValueError:
        relative = Path(source.name)
    destination = attachments_root / relative.parent / source.name
    if destination.exists():
        shutil.rmtree(destination)
    destination.mkdir(parents=True, exist_ok=True)
    return destination


def _write_zip_images(
    source: Path,
    attachments_root: Path,
    member_prefix: str,
    location_prefix: str,
) -> list[EvidenceAttachment]:
    destination = _attachment_dir(source, attachments_root)
    attachments: list[EvidenceAttachment] = []
    with zipfile.ZipFile(source) as archive:
        members = sorted(
            name for name in archive.namelist()
            if name.startswith(member_prefix) and not name.endswith("/")
        )
        for number, member in enumerate(members, 1):
            suffix = Path(member).suffix.lower() or ".bin"
            output = destination / f"image-{number:03d}{suffix}"
            output.write_bytes(archive.read(member))
            attachments.append(EvidenceAttachment(
                source_path=str(source),
                attachment_path=str(output),
                location=f"{location_prefix} {number}",
                media_type=mimetypes.guess_type(output.name)[0] or "application/octet-stream",
            ))
    return attachments


def _extract_text(path: Path) -> DocumentExtraction:
    return DocumentExtraction(text=path.read_text(encoding="utf-8", errors="replace"))


def _extract_pdf(path: Path, attachments_root: Path) -> DocumentExtraction:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DependencyUnavailableError("pypdf", "PDF") from exc

    reader = PdfReader(path)
    pages: list[str] = []
    warnings: list[str] = []
    attachments: list[EvidenceAttachment] = []
    destination = _attachment_dir(path, attachments_root)
    for page_number, page in enumerate(reader.pages, 1):
        text = page.extract_text() or ""
        pages.append(f"[PDF page {page_number}]\n{text}")
        if not text.strip():
            warnings.append(f"page {page_number}: no extractable text")
        try:
            images = list(page.images)
        except ImportError:
            warnings.append(f"page {page_number}: image extraction requires the 'Pillow' package")
            continue
        except Exception as exc:  # pypdf exposes unsupported image filters here.
            warnings.append(f"page {page_number}: image extraction failed: {exc}")
            continue
        for image_number, image in enumerate(images, 1):
            safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", image.name or "image.bin")
            output = destination / f"page-{page_number:04d}-image-{image_number:03d}-{safe_name}"
            output.write_bytes(image.data)
            attachments.append(EvidenceAttachment(
                source_path=str(path),
                attachment_path=str(output),
                location=f"page {page_number}, image {image_number}",
                media_type=mimetypes.guess_type(output.name)[0] or "application/octet-stream",
            ))
    missing = ["Pillow"] if any("'Pillow'" in warning for warning in warnings) else []
    return DocumentExtraction(
        text="\n\n".join(pages),
        attachments=attachments,
        warnings=warnings,
        missing_dependencies=missing,
    )


def _extract_docx(path: Path, attachments_root: Path) -> DocumentExtraction:
    try:
        from docx import Document
    except ImportError as exc:
        raise DependencyUnavailableError("python-docx", "DOCX") from exc

    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table_number, table in enumerate(document.tables, 1):
        blocks.append(f"[DOCX table {table_number}]")
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    attachments = _write_zip_images(path, attachments_root, "word/media/", "document image")
    return DocumentExtraction(text="\n".join(blocks), attachments=attachments)


def _extract_xlsx(path: Path, attachments_root: Path) -> DocumentExtraction:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DependencyUnavailableError("openpyxl", "XLSX") from exc

    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[str] = []
    try:
        for sheet in workbook.worksheets:
            blocks.append(f"[XLSX sheet {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(values):
                    blocks.append("\t".join(values))
    finally:
        workbook.close()
    attachments = _write_zip_images(path, attachments_root, "xl/media/", "workbook image")
    return DocumentExtraction(text="\n".join(blocks), attachments=attachments)


def extract_document(path: Path, attachments_root: Path) -> DocumentExtraction:
    """Extract text and image evidence without moving the source document."""
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return _extract_text(path)
    if suffix == ".pdf":
        return _extract_pdf(path, attachments_root)
    if suffix == ".docx":
        return _extract_docx(path, attachments_root)
    if suffix == ".xlsx":
        return _extract_xlsx(path, attachments_root)
    raise ValueError(f"unsupported document type: {suffix or '<none>'}")
